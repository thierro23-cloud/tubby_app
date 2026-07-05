# =============================================================================
# 🧾 SERVICIO: generar_informe_control_obras.py
# =============================================================================
#
# MOTOR DE INFORMES DE CONTROL DE OBRAS
#
# Este servicio es utilizado por:
#
#   • scripts automáticos
#   • blueprints Flask
#   • panel de estadísticas
#
# Permite:
#
#   1️⃣ Generar datos del informe
#   2️⃣ Generar informe mensual automático
#   3️⃣ Generar informe entre fechas
#   4️⃣ Generar DOCX
#   5️⃣ Generar PDF
#   6️⃣ Registrar auditoría
#
# =============================================================================
# ÍNDICE
# =============================================================================
#
# 1️⃣ IMPORTS
# 2️⃣ UTILIDADES DE FECHAS
# 3️⃣ CONFIGURACIÓN RUTA REPORTES
# 4️⃣ QUERY PRINCIPAL INFORME
# 5️⃣ AGRUPAR POR OBRA
# 6️⃣ GENERAR DATOS PARA FLASK
# 7️⃣ GENERAR DOCX
# 8️⃣ GENERAR PDF
# 9️⃣ REGISTRO EN BD
# 🔟 GENERAR INFORME ENTRE FECHAS
# 1️⃣1️⃣ GENERAR INFORME MENSUAL
# =============================================================================


# =============================================================================
# 1️⃣ IMPORTS
# =============================================================================

from __future__ import annotations

import os
from datetime import date, datetime
from typing import Dict, Any, List, Tuple

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

from db import ejecutar_query, ejecutar_non_query


# =============================================================================
# 2️⃣ CÁLCULO MES ANTERIOR
# =============================================================================

def rango_mes_anterior(hoy: date | None = None) -> Tuple[date, date, int, int]:

    hoy = hoy or date.today()

    year = hoy.year
    month = hoy.month

    if month == 1:
        prev_year = year - 1
        prev_month = 12
    else:
        prev_year = year
        prev_month = month - 1

    inicio = date(prev_year, prev_month, 1)
    fin = date(year, month, 1)

    return inicio, fin, prev_year, prev_month


# =============================================================================
# 3️⃣ RUTA DE REPORTES
# =============================================================================

def get_ruta_reportes():

    sql = """
    SELECT valor
    FROM bd_tbl_comunes.tbl_app_config
    WHERE clave='RUTA_REPORTES_CONTROL_OBRAS'
    LIMIT 1
    """

    filas = ejecutar_query(sql, (), nombre_bd="control_via_publica")

    if not filas:
        raise RuntimeError("No existe RUTA_REPORTES_CONTROL_OBRAS")

    return filas[0]["valor"]


# =============================================================================
# 4️⃣ CARGAR DATOS INFORME
# =============================================================================

def cargar_datos_informe(inicio, fin):

    sql = """
    SELECT
      o.idtbl_obras,
      o.observaciones,
      o.fecha_inicio,
      o.fecha_finalizacion,

      p.Nombre_Razon_Social proveedor_nombre,
      p.NIF proveedor_nif,

      c.calles calle_nombre,
      tv.tipos_de_vias tipo_via_nombre,

      cvp.idtbl_control_via_publica,
      cvp.fecha_inspeccion,
      cvp.vallas,
      cvp.vallas_metros,
      cvp.materiales_de_construccion,
      cvp.materiales_metros,
      cvp.andamios,
      cvp.gruas,
      cvp.n_agente1,
      cvp.idtbl_gestores gestor_control_id

    FROM control_via_publica.tbl_obras o

    LEFT JOIN bd_tbl_comunes.tbl_proveedores p
      ON p.Idtbl_proveedores = o.idtbl_proveedor

    LEFT JOIN bd_tbl_comunes.tbl_calles c
      ON c.idtbl_calles = o.idtbl_calles

    LEFT JOIN bd_tbl_comunes.tbl_tipos_de_vias tv
      ON tv.idtbl_tipos_de_vias = o.idtbl_tipos_de_vias

    LEFT JOIN control_via_publica.tbl_control_via_publica cvp
      ON cvp.idtbl_obras = o.idtbl_obras
     AND cvp.fecha_inspeccion >= %s
     AND cvp.fecha_inspeccion < %s

    WHERE (o.fecha_finalizacion IS NULL OR o.fecha_finalizacion='')

    ORDER BY o.idtbl_obras DESC, cvp.fecha_inspeccion
    """

    return ejecutar_query(sql, (inicio, fin), nombre_bd="control_via_publica")


# =============================================================================
# 5️⃣ AGRUPAR POR OBRA
# =============================================================================

def agrupar_por_obra(filas):

    out = {}

    for r in filas:

        oid = r["idtbl_obras"]

        if oid not in out:

            out[oid] = {
                "obra": {

                    "id": oid,
                    "proveedor": r.get("proveedor_nombre"),
                    "nif": r.get("proveedor_nif"),
                    "tipo_via": r.get("tipo_via_nombre"),
                    "calle": r.get("calle_nombre"),
                    "observaciones": r.get("observaciones")

                },
                "visitas": []
            }

        if r.get("idtbl_control_via_publica"):

            out[oid]["visitas"].append({

                "fecha": r.get("fecha_inspeccion"),
                "agente": r.get("n_agente1"),
                "vallas": r.get("vallas"),
                "metros_vallas": r.get("vallas_metros"),
                "materiales": r.get("materiales_de_construccion"),
                "metros_materiales": r.get("materiales_metros"),
                "andamios": r.get("andamios"),
                "gruas": r.get("gruas"),
                "gestor": r.get("gestor_control_id")

            })

    return out


# =============================================================================
# 6️⃣ GENERAR DATOS PARA FLASK
# =============================================================================

def generar_informe_datos(fecha_inicio=None, fecha_fin=None):

    filas = cargar_datos_informe(fecha_inicio, fecha_fin)

    return agrupar_por_obra(filas)


# =============================================================================
# 7️⃣ GENERAR DOCX
# =============================================================================

def generar_docx(path, datos, anio, mes):

    doc = Document()

    doc.add_heading(f"Informe Control Obras {anio}-{mes:02d}", level=1)

    for oid, pack in datos.items():

        obra = pack["obra"]
        visitas = pack["visitas"]

        doc.add_heading(f"Obra {oid}", level=2)

        doc.add_paragraph(f"Proveedor: {obra['proveedor']}")
        doc.add_paragraph(f"Ubicación: {obra['tipo_via']} {obra['calle']}")

        if not visitas:

            doc.add_paragraph("Sin visitas")

        else:

            for v in visitas:

                doc.add_paragraph(
                    f"{v['fecha']} | agente {v['agente']}"
                )

    os.makedirs(os.path.dirname(path), exist_ok=True)

    doc.save(path)


# =============================================================================
# 8️⃣ GENERAR PDF
# =============================================================================

def generar_pdf(path, datos, anio, mes):

    os.makedirs(os.path.dirname(path), exist_ok=True)

    c = canvas.Canvas(path, pagesize=A4)

    y = 800

    c.drawString(50, y, f"Informe Control Obras {anio}-{mes:02d}")

    y -= 40

    for oid, pack in datos.items():

        obra = pack["obra"]

        c.drawString(50, y, f"Obra {oid} - {obra['proveedor']}")

        y -= 20

        if y < 100:
            c.showPage()
            y = 800

    c.save()


# =============================================================================
# 9️⃣ REGISTRAR INFORME
# =============================================================================

def registrar_informe(anio, mes, docx, pdf, estado, error):

    sql = """
    INSERT INTO tbl_informes_control_obras
    (periodo_anio, periodo_mes, ruta_docx, ruta_pdf, estado, error_detalle)
    VALUES (%s,%s,%s,%s,%s,%s)
    """

    ejecutar_non_query(
        sql,
        (anio, mes, docx, pdf, estado, error),
        nombre_bd="control_via_publica"
    )


# =============================================================================
# 🔟 GENERAR INFORME ENTRE FECHAS (USADO POR FLASK)
# =============================================================================

def generar_informe_entre_fechas(fecha_inicio, fecha_fin):

    filas = cargar_datos_informe(fecha_inicio, fecha_fin)

    return agrupar_por_obra(filas)


# =============================================================================
# 1️⃣1️⃣ GENERAR INFORME AUTOMÁTICO
# =============================================================================

def generar_informe_control_obras():

    inicio, fin, anio, mes = rango_mes_anterior()

    carpeta = get_ruta_reportes()

    base = f"control_obras_{anio}_{mes:02d}"

    docx = os.path.join(carpeta, base + ".docx")
    pdf = os.path.join(carpeta, base + ".pdf")

    try:

        filas = cargar_datos_informe(inicio, fin)

        datos = agrupar_por_obra(filas)

        generar_docx(docx, datos, anio, mes)

        generar_pdf(pdf, datos, anio, mes)

        registrar_informe(anio, mes, docx, pdf, "ok", None)

    except Exception as e:

        registrar_informe(anio, mes, docx, pdf, "error", str(e))

        raise
# =========================================================
# 🚗 MÓDULO: RIO_TORIO - CONSULTAS AVANZADAS DE PLAZAS
# =========================================================
# COMIENZO MÓDULO
# ---------------------------------------------------------
# Este módulo implementa la lógica de backend para las
# "Consultas avanzadas" de plazas del parquin Río Torío.
#
# Incluye:
#   - Consulta del estado ACTUAL de plazas (ocupadas/libres)
#     con filtros dinámicos.
#   - Exportación del estado actual a Excel y Word.
#   - Consulta de BAJAS de plazas en un periodo (histórico).
#   - Exportación de BAJAS del periodo a Excel y Word.
#
# Estructura:
#   1. Configuración del Blueprint y constantes SQL base.
#   2. Funciones de exportación del estado actual (Excel/Word).
#   3. Funciones de consulta del histórico de bajas.
#   4. Funciones de exportación de bajas del periodo (Excel/Word).
#   5. Vista principal Flask: lectura de filtros, consultas,
#      exportaciones y renderizado de plantilla HTML.
# ---------------------------------------------------------


from flask import Blueprint, render_template, request, send_file
from db import get_connection
import io
import xlsxwriter
from docx import Document

# =========================================================
# 1. CONFIGURACIÓN DEL BLUEPRINT Y SQL BASE
# =========================================================
# COMIENZO SECCIÓN 1
# ---------------------------------------------------------
# Se define el Blueprint específico para las consultas avanzadas
# de Río Torío, junto con las consultas SQL base para:
#   - Estado actual de plazas + usuarios vinculados.
#   - Listado de proveedores para filtros.
#   - Consulta de proveedores (BD comunes) por IDs.
# ---------------------------------------------------------

btn_rio_torio_consultas_avanzadas_bp = Blueprint(
    "btn_rio_torio_consultas_avanzadas_bp",
    __name__,
    url_prefix="/parquin/rio_torio/consultas_avanzadas",
)

# 1.1. Consulta general de plazas + usuarios actuales
SQL_PLAZAS_Y_USUARIOS = """
    SELECT
        -- PLAZAS
        p.idtbl_plazas,
        p.codigo_plazas,
        p.observaciones,
        p.fila,
        p.numero_expediente,
        p.idtbl_usuarios,
        p.fecha_inicio      AS plaza_fecha_inicio,
        p.fecha_fin         AS plaza_fecha_fin,
        p.exp_solicitud     AS plaza_exp_solicitud,
        p.exp_solicitud_fin AS plaza_exp_solicitud_fin,
        p.idtbl_inventario,

        -- USUARIOS
        u.idtbl_usuarios          AS idtbl_gestores,
        u.idtbl_proveedores       AS idtbl_gestores_proveedor,
        u.numero_cuenta           AS usuario_numero_cuenta,
        u.activo_baja             AS usuario_activo_baja,
        u.fecha_inicio            AS usuario_fecha_inicio,
        u.fecha_baja              AS usuario_fecha_baja,
        u.rol                     AS usuario_rol

    FROM tbl_plazas AS p
    LEFT JOIN tbl_usuarios AS u
           ON u.idtbl_usuarios = p.idtbl_usuarios
    WHERE 1=1
"""

# 1.2. Listado de proveedores activos del parquin para filtros
SQL_PROVEEDORES_FILTRO = """
    SELECT
        Idtbl_proveedores,
        Nombre_Razon_Social,
        NIF
    FROM tbl_proveedores
    WHERE parquin = 1
    ORDER BY Nombre_Razon_Social
"""

# 1.3. Consulta de proveedores en BD comunes a partir de lista de IDs
SQL_PROVEEDORES_COMUNES = """
    SELECT
        Idtbl_proveedores,
        Nombre_Razon_Social,
        NIF
    FROM tbl_proveedores
    WHERE parquin = 1
      AND Idtbl_proveedores IN ({placeholders})
"""

# ---------------------------------------------------------
# FIN SECCIÓN 1
# =========================================================


# =========================================================
# 2. EXPORTACIÓN ESTADO ACTUAL (EXCEL / WORD)
# =========================================================
# COMIENZO SECCIÓN 2
# ---------------------------------------------------------
# Estas funciones exportan el ESTADO ACTUAL de las plazas (la
# búsqueda con filtros) a Excel y Word:
#
#   - _exportar_excel:
#       Genera un XLSX con:
#         · Título + fechas de consulta.
#         · Totales (libres/ocupadas/total).
#         · Tabla con el detalle de plazas.
#
#   - _exportar_word:
#       Genera un DOCX con:
#         · Título + fechas de consulta.
#         · Totales.
#         · Tabla con el detalle de plazas.
# ---------------------------------------------------------


# 2.1. Exportar estado actual a Excel
def _exportar_excel(plazas, libres, ocupadas, total, fecha_ini=None, fecha_fin=None):
    """
    Genera un Excel con el estado actual de las plazas, aplicando los filtros
    utilizados en la consulta (fila, estado, fechas, proveedor, etc.).

    Estructura de la hoja:
        - Fila 0: Título.
        - Fila 1: Fechas de consulta.
        - Fila 3: Totales (libres/ocupadas/total).
        - Fila 5: Cabeceras de columnas.
        - A partir de fila 6: Detalle de plazas.
    """
    output = io.BytesIO()
    workbook = xlsxwriter.Workbook(output, {"in_memory": True})
    ws = workbook.add_worksheet("Plazas")

    # Formatos básicos para título e información
    title_format = workbook.add_format({"bold": True, "font_size": 14})
    info_format = workbook.add_format({"italic": True})

    # 2.1.1 Título del listado
    ws.write(0, 0, "Listado de plazas de parking (estado actual)", title_format)

    # 2.1.2 Texto de periodo / fechas de consulta
    texto_periodo = "Fechas de consulta: "
    if fecha_ini and fecha_fin:
        texto_periodo += f"{fecha_ini} a {fecha_fin}"
    elif fecha_ini:
        texto_periodo += f"desde {fecha_ini}"
    elif fecha_fin:
        texto_periodo += f"hasta {fecha_fin}"
    else:
        texto_periodo += "no especificadas"
    ws.write(1, 0, texto_periodo, info_format)

    # 2.1.3 Totales de plazas
    ws.write(3, 0, "TOTALES")
    ws.write(3, 3, f"Libres: {libres}")
    ws.write(3, 4, f"Ocupadas: {ocupadas}")
    ws.write(3, 5, f"Total: {total}")

    # 2.1.4 Cabeceras de la tabla
    headers = [
        "ID plaza",
        "Código",
        "Fila",
        "Estado",
        "Proveedor",
        "NIF",
        "Alta plaza",
        "Baja plaza",
        "F.inicio usuario",
        "F.baja usuario",
        "Exp. alta plaza",
        "Exp. baja plaza",
    ]
    for col, h in enumerate(headers):
        ws.write(5, col, h)

    # 2.1.5 Datos de detalle de plazas
    for row, p in enumerate(plazas, start=6):
        ws.write(row, 0, p["idtbl_plazas"])
        ws.write(row, 1, p["codigo_plazas"])
        ws.write(row, 2, p["fila"])
        ws.write(row, 3, "Ocupada" if p["idtbl_usuarios"] else "Libre")
        ws.write(row, 4, p["prov_nombre"] or "")
        ws.write(row, 5, p["prov_nif"] or "")
        ws.write(row, 6, str(p["plaza_fecha_inicio"] or ""))
        ws.write(row, 7, str(p["plaza_fecha_fin"] or ""))
        ws.write(row, 8, str(p["usuario_fecha_inicio"] or ""))
        ws.write(row, 9, str(p["usuario_fecha_baja"] or ""))
        ws.write(row, 10, str(p["plaza_exp_solicitud"] or ""))
        ws.write(row, 11, str(p["plaza_exp_solicitud_fin"] or ""))

    workbook.close()
    output.seek(0)
    return send_file(
        output,
        as_attachment=True,
        download_name="plazas_parquin_detalle.xlsx",
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )


# 2.2. Exportar estado actual a Word
def _exportar_word(plazas, libres, ocupadas, total, fecha_ini=None, fecha_fin=None):
    """
    Genera un documento Word con el estado actual de las plazas.

    Estructura del documento:
        - Título principal.
        - Párrafo con fechas de consulta.
        - Párrafo con totales (libres/ocupadas/total).
        - Tabla de detalle de plazas.
    """
    doc = Document()
    doc.add_heading("Listado de plazas de parking (detalle)", level=1)

    # 2.2.1 Fechas de consulta
    p_fechas = doc.add_paragraph()
    texto_periodo = "Fechas de consulta: "
    if fecha_ini and fecha_fin:
        texto_periodo += f"{fecha_ini} a {fecha_fin}"
    elif fecha_ini:
        texto_periodo += f"desde {fecha_ini}"
    elif fecha_fin:
        texto_periodo += f"hasta {fecha_fin}"
    else:
        texto_periodo += "no especificadas"
    p_fechas.add_run(texto_periodo).italic = True

    # 2.2.2 Totales
    p_info = doc.add_paragraph()
    p_info.add_run(
        f"Total: {total}  ·  Libres: {libres}  ·  Ocupadas: {ocupadas}"
    ).bold = True

    # Separación visual antes de la tabla
    doc.add_paragraph()

    # 2.2.3 Tabla de detalle de plazas
    table = doc.add_table(rows=1, cols=11)
    hdr = table.rows[0].cells
    hdr[0].text = "Plaza"
    hdr[1].text = "Fila"
    hdr[2].text = "Estado"
    hdr[3].text = "Proveedor"
    hdr[4].text = "NIF"
    hdr[5].text = "Alta plaza"
    hdr[6].text = "Baja plaza"
    hdr[7].text = "F.inicio usuario"
    hdr[8].text = "F.baja usuario"
    hdr[9].text = "Exp. alta plaza"
    hdr[10].text = "Exp. baja plaza"

    for p in plazas:
        row = table.add_row().cells
        row[0].text = f"{p['codigo_plazas'] or ''} (ID {p['idtbl_plazas']})"
        row[1].text = str(p["fila"] or "")
        row[2].text = "Ocupada" if p["idtbl_usuarios"] else "Libre"
        row[3].text = p["prov_nombre"] or ""
        row[4].text = p["prov_nif"] or ""
        row[5].text = str(p["plaza_fecha_inicio"] or "")
        row[6].text = str(p["plaza_fecha_fin"] or "")
        row[7].text = str(p["usuario_fecha_inicio"] or "")
        row[8].text = str(p["usuario_fecha_baja"] or "")
        row[9].text = str(p["plaza_exp_solicitud"] or "")
        row[10].text = str(p["plaza_exp_solicitud_fin"] or "")

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return send_file(
        output,
        as_attachment=True,
        download_name="plazas_parquin_detalle.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


# ---------------------------------------------------------
# FIN SECCIÓN 2
# =========================================================


# =========================================================
# 3. CONSULTA DEL HISTÓRICO: BAJAS EN PERIODO
# =========================================================
# COMIENZO SECCIÓN 3
# ---------------------------------------------------------
# Esta sección contiene la función para recuperar las BAJAS
# de plazas en un periodo (histórico):
#
#   - obtener_bajas_plazas_periodo:
#       Devuelve plazas cuyo registro en tbl_historico_plazas
#       tiene fecha_fin no nula y comprendida en el rango.
# ---------------------------------------------------------


def obtener_bajas_plazas_periodo(hist_fecha_ini, hist_fecha_fin):
    """
    Devuelve una lista de dicts con las plazas que han tenido una BAJA
    (fecha_fin no nula) dentro del rango [hist_fecha_ini, hist_fecha_fin].

    Criterio:
        - h.fecha_fin IS NOT NULL
        - h.fecha_fin BETWEEN hist_fecha_ini AND hist_fecha_fin.
    """
    if not hist_fecha_ini or not hist_fecha_fin:
        return []

    sql = """
        SELECT
            p.codigo_plazas,
            p.fila,
            p.numero_expediente,
            h.fecha_inicio,
            h.fecha_fin,
            h.exp_solicitud_fin,
            pr.Nombre_Razon_Social,
            pr.NIF
        FROM parquin_camiones.tbl_historico_plazas AS h
        INNER JOIN parquin_camiones.tbl_plazas AS p
            ON p.idtbl_plazas = h.idtbl_plazas
        INNER JOIN bd_tbl_comunes.tbl_proveedores AS pr
            ON pr.Idtbl_proveedores = h.idtbl_proveedores
        WHERE
            h.fecha_fin IS NOT NULL
            AND h.fecha_fin BETWEEN %s AND %s
        ORDER BY p.fila, p.codigo_plazas, h.fecha_fin
    """

    conn = get_connection("parquin_camiones")
    cursor = conn.cursor(dictionary=True)
    params = (hist_fecha_ini, hist_fecha_fin)
    cursor.execute(sql, params)
    filas = cursor.fetchall()
    cursor.close()
    conn.close()

    return filas


# ---------------------------------------------------------
# FIN SECCIÓN 3
# =========================================================


# =========================================================
# 4. EXPORTACIÓN BAJAS DEL PERIODO (EXCEL / WORD)
# =========================================================
# COMIENZO SECCIÓN 4
# ---------------------------------------------------------
# Estas funciones exportan SOLO las BAJAS del periodo:
#
#   - _exportar_periodo_excel_bajas:
#       Genera un XLSX con:
#         · Título + texto de periodo.
#         · Tabla única de bajas.
#
#   - _exportar_periodo_word_bajas:
#       Genera un DOCX con:
#         · Título + periodo.
#         · Sección "Bajas en el periodo" + tabla.
# ---------------------------------------------------------


# 4.1. Exportar BAJAS del periodo a Excel
def _exportar_periodo_excel_bajas(bajas_periodo, hist_fecha_ini, hist_fecha_fin):
    """
    Genera un Excel con SOLO las BAJAS de plazas en el periodo.
    """
    output = io.BytesIO()
    workbook = xlsxwriter.Workbook(output, {"in_memory": True})

    title_format = workbook.add_format({"bold": True, "font_size": 14})
    info_format = workbook.add_format({"italic": True})

    headers = [
        "Plaza",
        "Fila",
        "Abonado",
        "NIF",
        "Fecha inicio",
        "Fecha fin",
        "Exp. fin",
        "Nº expediente plaza",
    ]

    ws = workbook.add_worksheet("Bajas periodo")

    ws.write(0, 0, "Bajas de plazas en el periodo", title_format)
    ws.write(
        1,
        0,
        f"Bajas con fecha_fin entre {hist_fecha_ini} y {hist_fecha_fin}",
        info_format,
    )

    for col, h in enumerate(headers):
        ws.write(3, col, h)

    for row, b in enumerate(bajas_periodo, start=4):
        ws.write(row, 0, b.get("codigo_plazas", "") or "")
        ws.write(row, 1, b.get("fila", "") or "")
        ws.write(row, 2, b.get("Nombre_Razon_Social", "") or "")
        ws.write(row, 3, b.get("NIF", "") or "")
        ws.write(row, 4, str(b.get("fecha_inicio", "") or ""))
        ws.write(row, 5, str(b.get("fecha_fin", "") or ""))
        ws.write(row, 6, b.get("exp_solicitud_fin", "") or "")
        ws.write(row, 7, b.get("numero_expediente", "") or "")

    workbook.close()
    output.seek(0)

    nombre = f"bajas_plazas_{hist_fecha_ini}_a_{hist_fecha_fin}.xlsx"
    return send_file(
        output,
        as_attachment=True,
        download_name=nombre,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )


# 4.2. Exportar BAJAS del periodo a Word
def _exportar_periodo_word_bajas(bajas_periodo, hist_fecha_ini, hist_fecha_fin):
    """
    Genera un Word con SOLO las BAJAS de plazas en el periodo.
    """
    doc = Document()
    doc.add_heading("Bajas de plazas de parking", level=1)

    p_info = doc.add_paragraph()
    p_info.add_run(f"Periodo: {hist_fecha_ini} a {hist_fecha_fin}").bold = True

    doc.add_paragraph()
    doc.add_heading("Bajas en el periodo", level=2)

    if bajas_periodo:
        table_b = doc.add_table(rows=1, cols=8)
        hdr_b = table_b.rows[0].cells
        hdr_b[0].text = "Plaza"
        hdr_b[1].text = "Fila"
        hdr_b[2].text = "Abonado"
        hdr_b[3].text = "NIF"
        hdr_b[4].text = "Fecha inicio"
        hdr_b[5].text = "Fecha fin"
        hdr_b[6].text = "Exp. fin"
        hdr_b[7].text = "Nº expediente"

        for b in bajas_periodo:
            row = table_b.add_row().cells
            row[0].text = str(b.get("codigo_plazas", "") or "")
            row[1].text = str(b.get("fila", "") or "")
            row[2].text = str(b.get("Nombre_Razon_Social", "") or "")
            row[3].text = str(b.get("NIF", "") or "")
            row[4].text = str(b.get("fecha_inicio", "") or "")
            row[5].text = str(b.get("fecha_fin", "") or "")
            row[6].text = str(b.get("exp_solicitud_fin", "") or "")
            row[7].text = str(b.get("numero_expediente", "") or "")
    else:
        doc.add_paragraph("No hay bajas en el periodo indicado.").italic = True

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)

    nombre = f"bajas_plazas_{hist_fecha_ini}_a_{hist_fecha_fin}.docx"
    return send_file(
        output,
        as_attachment=True,
        download_name=nombre,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


# ---------------------------------------------------------
# FIN SECCIÓN 4
# =========================================================


# =========================================================
# 5. VISTA PRINCIPAL: CONSULTAS AVANZADAS (ENDPOINT)
# =========================================================
# COMIENZO SECCIÓN 5
# ---------------------------------------------------------
# Esta vista:
#   - Lee los filtros de estado actual e histórico.
#   - Consulta plazas + usuarios según filtros.
#   - Enriquede los datos con proveedores (BD comunes).
#   - Aplica filtros en memoria por nombre/NIF proveedor.
#   - Calcula totales de plazas (libres/ocupadas).
#   - Carga listado de proveedores para desplegables.
#   - Gestiona exportaciones:
#       · exportar=excel/word → estado actual.
#       · exportar_periodo=excel/word → bajas del periodo.
#   - Si no se exporta, renderiza la plantilla HTML.
# ---------------------------------------------------------


@btn_rio_torio_consultas_avanzadas_bp.route(
    "/", methods=["GET", "POST"], endpoint="btn_rio_torio_consultas_avanzadas"
)
def btn_rio_torio_consultas_avanzadas():
    """Vista principal de consultas avanzadas de plazas (Río Torío)."""

    # 5.1. Lectura de parámetros de filtro (estado actual)
    fila = request.values.get("fila")
    estado = request.values.get("estado")
    codigo = request.values.get("codigo")
    prov_nombre = request.values.get("prov_nombre")
    prov_nif = request.values.get("prov_nif")
    fecha_ini = request.values.get("fecha_ini")
    fecha_fin = request.values.get("fecha_fin")
    exportar = request.values.get("exportar")

    # 5.2. Lectura de parámetros (histórico por periodo)
    hist_fecha_ini = request.values.get("hist_fecha_ini")
    hist_fecha_fin = request.values.get("hist_fecha_fin")
    exportar_periodo = request.values.get("exportar_periodo")

    # 5.3. Consulta principal: plazas + usuarios actuales (aplicando filtros)
    conn = get_connection("parquin_camiones")
    cursor = conn.cursor(dictionary=True)
    sql = SQL_PLAZAS_Y_USUARIOS
    params = []

    # Filtro por fila
    if fila:
        sql += " AND p.fila = %s"
        params.append(fila)

    # Filtro por estado (libre/ocupada)
    if estado:
        e = estado.strip().lower()
        if e == "libre":
            sql += " AND p.idtbl_usuarios IS NULL"
        elif e == "ocupada":
            sql += " AND p.idtbl_usuarios IS NOT NULL"

    # Filtro por código de plaza (LIKE)
    if codigo:
        sql += " AND p.codigo_plazas LIKE %s"
        params.append(f"%{codigo}%")

    # Filtros por fechas de plaza (alta/baja)
    if fecha_ini:
        sql += " AND (p.fecha_inicio IS NULL OR p.fecha_inicio >= %s)"
        params.append(fecha_ini)

    if fecha_fin:
        sql += " AND (p.fecha_fin IS NULL OR p.fecha_fin <= %s)"
        params.append(fecha_fin)

    # Ordenación por fila y código
    sql += " ORDER BY p.fila, p.codigo_plazas"
    cursor.execute(sql, tuple(params))
    plazas = cursor.fetchall()
    cursor.close()
    conn.close()

    # 5.4. Obtención de proveedores relacionados (BD comunes)
    ids_proveedores = {
        int(p["idtbl_gestores_proveedor"])
        for p in plazas
        if p.get("idtbl_gestores_proveedor") is not None
    }

    proveedor_por_id = {}
    if ids_proveedores:
        conn_comunes = get_connection("bd_tbl_comunes")
        cursor_comunes = conn_comunes.cursor(dictionary=True)
        placeholders = ",".join(["%s"] * len(ids_proveedores))
        sql_prov = SQL_PROVEEDORES_COMUNES.format(placeholders=placeholders)
        cursor_comunes.execute(sql_prov, tuple(ids_proveedores))

        for row in cursor_comunes.fetchall():
            prov_id = int(row["Idtbl_proveedores"])
            proveedor_por_id[prov_id] = {
                "prov_id": prov_id,
                "prov_nombre": row["Nombre_Razon_Social"],
                "prov_nif": row["NIF"],
            }

        cursor_comunes.close()
        conn_comunes.close()

    # 5.5. Enriquecimiento de las plazas con datos de proveedor
    plazas_enriquecidas = []
    for p in plazas:
        uid = p.get("idtbl_gestores_proveedor")
        prov = proveedor_por_id.get(int(uid)) if uid is not None else None
        p["prov_id"] = prov["prov_id"] if prov else None
        p["prov_nombre"] = prov["prov_nombre"] if prov else None
        p["prov_nif"] = prov["prov_nif"] if prov else None
        plazas_enriquecidas.append(p)
    plazas = plazas_enriquecidas

    # 5.6. Filtros adicionales en memoria (proveedor nombre / NIF)
    if prov_nombre:
        t = prov_nombre.lower().strip()
        plazas = [
            p for p in plazas if p["prov_nombre"] and t in p["prov_nombre"].lower()
        ]

    if prov_nif:
        t = prov_nif.lower().strip()
        plazas = [p for p in plazas if p["prov_nif"] and t in p["prov_nif"].lower()]

    # 5.7. Cálculo de totales (libres/ocupadas/total)
    total = len(plazas)
    ocupadas = sum(1 for p in plazas if p["idtbl_usuarios"] is not None)
    libres = total - ocupadas

    # 5.8. Listado de proveedores para el desplegable de filtros
    conn_comunes = get_connection("bd_tbl_comunes")
    cursor_comunes = conn_comunes.cursor(dictionary=True)
    cursor_comunes.execute(SQL_PROVEEDORES_FILTRO)
    proveedores_filtro = cursor_comunes.fetchall()
    cursor_comunes.close()
    conn_comunes.close()

    # =====================================================
    # 5.9. EXPORTACIONES: ESTADO ACTUAL (EXCEL / WORD)
    # =====================================================
    if exportar == "excel":
        return _exportar_excel(plazas, libres, ocupadas, total, fecha_ini, fecha_fin)
    if exportar == "word":
        return _exportar_word(plazas, libres, ocupadas, total, fecha_ini, fecha_fin)

    # =====================================================
    # 5.10. EXPORTACIONES: BAJAS DEL PERIODO (HISTÓRICO)
    # =====================================================
    if exportar_periodo in ("excel", "word") and hist_fecha_ini and hist_fecha_fin:
        bajas_periodo = obtener_bajas_plazas_periodo(hist_fecha_ini, hist_fecha_fin)

        if exportar_periodo == "excel":
            return _exportar_periodo_excel_bajas(
                bajas_periodo, hist_fecha_ini, hist_fecha_fin
            )
        else:
            return _exportar_periodo_word_bajas(
                bajas_periodo, hist_fecha_ini, hist_fecha_fin
            )

    # 5.11. Si no hay exportación, se consulta bajas_periodo para mostrar en pantalla
    bajas_periodo = []
    if hist_fecha_ini and hist_fecha_fin:
        bajas_periodo = obtener_bajas_plazas_periodo(hist_fecha_ini, hist_fecha_fin)

    # 5.12. Renderizado de la plantilla HTML
    return render_template(
        "parquin/RIO_TORIO/consultas_avanzadas.html",
        plazas=plazas,
        fila=fila,
        estado=estado,
        codigo=codigo,
        prov_nombre=prov_nombre,
        prov_nif=prov_nif,
        fecha_ini=fecha_ini,
        fecha_fin=fecha_fin,
        hist_fecha_ini=hist_fecha_ini,
        hist_fecha_fin=hist_fecha_fin,
        total=total,
        libres=libres,
        ocupadas=ocupadas,
        proveedores_filtro=proveedores_filtro,
        bajas_periodo=bajas_periodo,
    )


# ---------------------------------------------------------
# FIN SECCIÓN 5
# =========================================================

# ---------------------------------------------------------
# FIN MÓDULO
# =========================================================

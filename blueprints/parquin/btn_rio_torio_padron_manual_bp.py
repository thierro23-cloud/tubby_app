# =============================================================================
# 🧾 PADRÓN RÍO TORÍO · HISTÓRICO COMO FUENTE ÚNICA (MANUAL)
# =============================================================================
"""
Botón de super admin para generar el padrón de Río Torío mes/año.

- Blueprint de botón:
    btn_rio_torio_padron_manual_bp
- Rutas:
    - /parquin/rio_torio/btn_rio_torio_padron_manual
        · GET  → muestra formulario con MES/AÑO
        · POST → valida MES/AÑO, genera DOCX y lo devuelve

- Fuente única de datos:
    - parquin_camiones.tbl_historico_plazas  (histórico de ocupación)
    - parquin_camiones.tbl_plazas           (plazas físicas)
    - parquin_camiones.tbl_usuarios         (usuarios de parking)
    - bd_tbl_comunes.tbl_proveedores        (proveedores / abonados)

Requisitos de acceso:
    - Usuario autenticado.
    - Rol: "gestor" o "super_admin".
"""

from __future__ import annotations

from datetime import date, timedelta
from collections import defaultdict
from pathlib import Path
import calendar

from flask import (
    Blueprint,
    current_app,
    flash,
    redirect,
    url_for,
    request,
    render_template,
    send_file,
)

from services.helpers import login_required, rol_required
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from db import ejecutar_query, get_connection

# Solo usamos print para depuración en import; NUNCA current_app a nivel de módulo.
print(">>> Cargando btn_rio_torio_padron_manual_bp.py")

# =============================================================================
# 1️⃣ LOCALIZAR CARPETA PADRÓN Y NOMBRES DE MESES
# =============================================================================

# BASE_DIR = raíz del proyecto (dos niveles por encima de este archivo)
BASE_DIR = Path(__file__).resolve().parent.parent


def localizar_padron(base_dir: Path) -> Path | None:
    """
    Busca recursivamente un directorio llamado 'padron' a partir de base_dir.
    Si lo encuentra, devuelve la primera ruta; si no, devuelve None.
    """
    for p in base_dir.rglob("padron"):
        if p.is_dir():
            return p
    return None


PADRON_DIR = localizar_padron(BASE_DIR)
if PADRON_DIR is None:
    # Carpeta por defecto si no existe ninguna 'padron' en el árbol
    PADRON_DIR = BASE_DIR / "parquin" / "padron"

RUTA_DESTINO = PADRON_DIR

# Nombres de meses en español, índice 1–12
MESES_ES = [
    "",
    "enero",
    "febrero",
    "marzo",
    "abril",
    "mayo",
    "junio",
    "julio",
    "agosto",
    "septiembre",
    "octubre",
    "noviembre",
    "diciembre",
]


# =============================================================================
# 2️⃣ BLUEPRINT (SIGUE CONVENCIÓN DEL PANEL PARQUIN)
#   - modulo:  modulo_parquin_rio_torio_bp  → modulo_id = "parquin_rio_torio"
#   - botones: btn_rio_torio_*_bp / btn_rio_torio_*
# =============================================================================

btn_rio_torio_padron_manual_bp = Blueprint(
    "btn_rio_torio_padron_manual_bp",
    __name__,
    url_prefix="/parquin/rio_torio",
)


# =============================================================================
# 3️⃣ HELPER: CÁLCULO DEL PERIODO DEL PADRÓN (JSON)
# =============================================================================

def obtener_periodo_padron(fecha_base: date | None = None) -> dict:
    """
    Calcula el periodo del padrón técnico.

    Regla:
      - Mes del padrón = mes anterior al primer día laborable del mes siguiente.

    Ejemplo:
      - Hoy 10/06/2026:
          · Mes siguiente: julio 2026.
          · Primer laborable de julio.
          · Mes del padrón: junio 2026 (1–30).

    Devuelve:
      {
        "inicio_mes": date,
        "fin_mes": date,
        "anio": int,
        "mes_num": int,
        "mes_nombre": str (nombre en inglés),
        "primer_laborable_mes_siguiente": date,
      }
    """
    if fecha_base is None:
        fecha_base = date.today()

    # Mes siguiente a fecha_base
    if fecha_base.month == 12:
        anio_siguiente = fecha_base.year + 1
        mes_siguiente = 1
    else:
        anio_siguiente = fecha_base.year
        mes_siguiente = fecha_base.month + 1

    # Primer día del mes siguiente
    primer_dia_mes_siguiente = date(anio_siguiente, mes_siguiente, 1)

    # Primer laborable (lunes–viernes) del mes siguiente
    d = primer_dia_mes_siguiente
    while d.weekday() >= 5:  # 5 = sábado, 6 = domingo
        d += timedelta(days=1)
    primer_laborable = d

    # Mes del padrón = mes anterior al mes_siguiente
    if mes_siguiente == 1:
        anio_padron = anio_siguiente - 1
        mes_padron = 12
    else:
        anio_padron = anio_siguiente
        mes_padron = mes_siguiente - 1

    # Inicio y fin del mes del padrón
    inicio_mes = date(anio_padron, mes_padron, 1)
    _, ultimo_dia = calendar.monthrange(anio_padron, mes_padron)
    fin_mes = date(anio_padron, mes_padron, ultimo_dia)

    mes_nombre = calendar.month_name[mes_padron]

    return {
        "inicio_mes": inicio_mes,
        "fin_mes": fin_mes,
        "anio": anio_padron,
        "mes_num": mes_padron,
        "mes_nombre": mes_nombre,
        "primer_laborable_mes_siguiente": primer_laborable,
    }


# =============================================================================
# 4️⃣ CONSULTA ÚNICA: HISTÓRICO + PLAZAS + USUARIOS + PROVEEDORES
# =============================================================================

def obtener_historico_periodo(inicio_mes: date, fin_mes: date) -> list[dict]:
    sql = (
        "SELECT "
        "    h.idtbl_historico_plazas, "
        "    h.idtbl_plazas, "
        "    h.idtbl_proveedores, "
        "    h.fecha_inicio, "
        "    h.fecha_fin, "
        "    h.exp_solicitud_fin, "
        "    h.forma_pago, "
        "    pl.codigo_plazas, "
        "    pl.fila, "
        "    CONCAT_WS(' ', p.apellidos, p.nombre) AS nombre_proveedor, "
        "    p.NIF AS nif_proveedor "
        "FROM parquin_camiones.tbl_historico_plazas AS h "
        "JOIN parquin_camiones.tbl_plazas AS pl "
        "  ON h.idtbl_plazas = pl.idtbl_plazas "
        "JOIN bd_tbl_comunes.tbl_proveedores AS p "
        "  ON h.idtbl_proveedores = p.Idtbl_proveedores "
        "WHERE "
        "    h.fecha_inicio <= %s "
        "    AND (h.fecha_fin IS NULL OR h.fecha_fin >= %s) "
        "ORDER BY p.apellidos, p.nombre, pl.codigo_plazas, h.fecha_inicio"
    )

    filas = ejecutar_query(
        sql,
        params=(fin_mes, inicio_mes),
        nombre_bd="parquin_camiones",
    )
    return filas
#=============================================================================
# 5️⃣ PADRÓN PRINCIPAL (AGREGADO POR PROVEEDOR)
# =============================================================================

def construir_padron_principal(filas: list[dict]) -> list[dict]:
    por_proveedor: dict[int, dict] = {}

    for f in filas:
        id_prov = f["idtbl_proveedores"]
        if id_prov not in por_proveedor:
            por_proveedor[id_prov] = {
                "nombre_proveedor": f["nombre_proveedor"],
                "nif": f["nif_proveedor"],
                "plazas": set(),
                "forma_pago": (f.get("forma_pago") 
                               or f.get("forma_pago_proveedor") 
                               or f.get("forma_pago_usuario") 
                               or ""),
            }
        por_proveedor[id_prov]["plazas"].add(f["codigo_plazas"])

    resultado = []
    for prov in por_proveedor.values():
        plazas_ordenadas = sorted(prov["plazas"], key=str)
        resultado.append(
            {
                "nombre_proveedor": prov["nombre_proveedor"],
                "nif": prov["nif"],
                "plazas": plazas_ordenadas,
                "total_plazas": len(plazas_ordenadas),
                "forma_pago": prov["forma_pago"],
            }
        )

    return resultado

# =============================================================================
# 6️⃣ VARIACIONES: BAJA, CAMBIO DE PLAZA, FORMA DE PAGO
# =============================================================================

def construir_variaciones(
    filas: list[dict],
    inicio_mes: date,
    fin_mes: date,
) -> list[dict]:
    """
    Construye la tabla de variaciones para el periodo.

    Tipos de cambio:
      - "BAJA":
          · fecha_fin en el mes del padrón.
          · Sin alta al día siguiente en otra plaza del mismo proveedor.
      - "Cambio de plaza":
          · Baja con fecha_fin = D.
          · Alta con fecha_inicio = D+1 en otra plaza del mismo proveedor.
      - "Forma de pago":
          · Cambio de forma_pago dentro del mes para la misma plaza.
    """
    variaciones: list[dict] = []

    # Agrupar registros por proveedor
    por_prov: dict[int, list[dict]] = defaultdict(list)
    for f in filas:
        por_prov[f["idtbl_proveedores"]].append(f)

    for id_prov, registros in por_prov.items():
        # Ordenar por fecha_inicio e id de plaza para facilitar el análisis
        registros.sort(key=lambda r: (r["fecha_inicio"], r["idtbl_plazas"]))

        bajas: list[dict] = []
        altas: list[dict] = []

        # Clasificar altas y bajas dentro del mes
        for r in registros:
            fi = r["fecha_inicio"]
            ff = r["fecha_fin"]

            if inicio_mes <= fi <= fin_mes:
                altas.append(r)

            if ff is not None and inicio_mes <= ff <= fin_mes:
                bajas.append(r)

        usadas_bajas: set[int] = set()
        usadas_altas: set[int] = set()

        # --- 1) Cambios de plaza: baja en D, alta en D+1 en otra plaza DEL MISMO PROVEEDOR ---
        for baja in bajas:
            ff = baja["fecha_fin"]
            if ff is None:
                continue

            dia_siguiente = ff + timedelta(days=1)

            for idx_alta, alta in enumerate(altas):
                if idx_alta in usadas_altas:
                    continue

                # Misma fecha (día siguiente) y mismo proveedor, pero plaza distinta
                if (
                    alta["fecha_inicio"] == dia_siguiente
                    and alta["idtbl_plazas"] != baja["idtbl_plazas"]
                    and alta["idtbl_proveedores"] == id_prov
                ):
                    variaciones.append(
                        {
                            "nombre_proveedor": baja["nombre_proveedor"],
                            "nif": baja["nif_proveedor"],
                            "numero_plaza": (
            f"{baja['codigo_plazas']} -> {alta['codigo_plazas']}"
        ),
                            "tipo_cambio": "Cambio de plaza",
                            "fecha": ff,
                            "forma_pago": baja.get("forma_pago") or baja.get("forma_pago_usuario"),
                        }
                    )
                    usadas_bajas.add(baja["idtbl_historico_plazas"])
                    usadas_altas.add(idx_alta)
                    break

        # --- 2) Bajas puras: sin alta al día siguiente ---
        for baja in bajas:
            if baja["idtbl_historico_plazas"] in usadas_bajas:
                continue

            variaciones.append(
                {
                    "nombre_proveedor": baja["nombre_proveedor"],
                    "nif": baja["nif_proveedor"],
                    "numero_plaza": baja["codigo_plazas"],
                    "tipo_cambio": "BAJA",
                    "fecha": baja["fecha_fin"],
                    "forma_pago": baja.get("forma_pago") or baja.get("forma_pago_usuario"),
                }
            )

        # --- 3) Cambios de forma de pago por plaza ---
        por_plaza: dict[int, list[dict]] = defaultdict(list)
        for r in registros:
            por_plaza[r["idtbl_plazas"]].append(r)

        for id_plaza, regs_plaza in por_plaza.items():
            regs_plaza.sort(key=lambda r: r["fecha_inicio"])
            forma_anterior = None

            for r in regs_plaza:
                forma_actual = r.get("forma_pago") or r.get("forma_pago_usuario")
                fi = r["fecha_inicio"]

                if forma_anterior is not None and forma_actual != forma_anterior:
                    if inicio_mes <= fi <= fin_mes:
                        variaciones.append(
                            {
                                "nombre_proveedor": r["nombre_proveedor"],
                                "nif": r["nif_proveedor"],
                                "numero_plaza": r["codigo_plazas"],
                                "tipo_cambio": "Forma de pago",
                                "fecha": fi,
                                "forma_pago": forma_actual,
                            }
                        )

                forma_anterior = forma_actual

    # Orden final para el informe
    variaciones.sort(
        key=lambda v: (v["nombre_proveedor"], v["fecha"], v["numero_plaza"])
    )
    return variaciones

# =============================================================================
# 7️⃣ RESUMEN PLAZAS (TOTAL / LIBRES / OCUPADAS)
# =============================================================================

def _obtener_resumen_plazas_rio_torio(inicio_mes: date, fin_mes: date) -> dict:
    """
    Obtiene resumen de plazas totales/libres/ocupadas en Río Torío.
    """
    sql = (
        "SELECT "
        "    COUNT(DISTINCT h.idtbl_plazas) AS total, "
        "    0 AS libres, "
        "    COUNT(DISTINCT h.idtbl_plazas) AS ocupadas "
        "FROM parquin_camiones.tbl_historico_plazas AS h "
        "WHERE "
        "    h.fecha_inicio <= %s "
        "    AND (h.fecha_fin IS NULL OR h.fecha_fin >= %s)"
    )

    filas = ejecutar_query(
        sql,
        params=(fin_mes, inicio_mes),
        nombre_bd="parquin_camiones",
    )
    row = filas[0] if filas else {"total": 0, "libres": 0, "ocupadas": 0}

    return {
        "totales": row["total"] or 0,
        "libres": row["libres"] or 0,
        "ocupadas": row["ocupadas"] or 0,
    }


# =============================================================================
# 8️⃣ GENERAR INFORME WORD · PADRÓN MENSUAL (CON CONTADOR POR TIPO DE PAGO)
# =============================================================================

def _generar_informe_word(mes_informe: int, anio_informe: int) -> Path:
    """
    Genera el Word del padrón mensual Río Torío.

    CARACTERÍSTICAS:
    -----------------
    - Contador por tipo de pago en el inicio (ANTES de la tabla):
        · total pago mensual    = X  (nº de plazas con forma_pago = 'mensual')
        · total pago trimestral = Y  (nº de plazas con forma_pago = 'trimestral')
        · total pago semestral  = Z  (nº de plazas con forma_pago = 'semestral')
        · total pago anual      = W  (nº de plazas con forma_pago = 'anual')
    - Una sola tabla principal con 5 columnas:
        · Apellidos y nombre
        · NIF
        · Forma de pago
        · Plazas (lista: 099, 074, ...)
        · Total plazas
    - Si hay variaciones, segunda tabla de variaciones.
    """

    # 1) Cálculo de periodo (primer y último día del mes)
    fecha_inicio_periodo = date(anio_informe, mes_informe, 1)
    if mes_informe == 12:
        fecha_inicio_mes_siguiente = date(anio_informe + 1, 1, 1)
    else:
        fecha_inicio_mes_siguiente = date(anio_informe, mes_informe + 1, 1)
    fecha_fin_periodo = fecha_inicio_mes_siguiente - timedelta(days=1)

    nombre_mes = MESES_ES[mes_informe].capitalize()
    nombre_fichero = (
        f"{nombre_mes}_{anio_informe}_Padron_de_parking_Rio_Torio.docx"
    )

    RUTA_DESTINO.mkdir(parents=True, exist_ok=True)
    ruta_completa = RUTA_DESTINO / nombre_fichero

    # 2) Obtener datos del histórico del periodo
    filas = obtener_historico_periodo(
        fecha_inicio_periodo,
        fecha_fin_periodo,
    )

    # 3) Construir padrón principal y variaciones
    padron_principal = construir_padron_principal(filas)
    variaciones = construir_variaciones(
        filas,
        fecha_inicio_periodo,
        fecha_fin_periodo,
    )

    # 4) Contadores por tipo de pago (forma_pago)
    #    Se calculan directamente sobre las filas del histórico
    #    Contamos plazas distintas por cada tipo de forma de pago.
    plazas_mensual = set()
    plazas_trimestral = set()
    plazas_semestral = set()
    plazas_anual = set()
    plazas_otro = set()

    for f in filas:
        # Preferimos forma_pago del histórico; si no, la del usuario
        forma_raw = (f.get("forma_pago") or f.get("forma_pago_usuario") or "").strip().lower()
        plaza = f.get("codigo_plazas")

        # Si no hay plaza, no tiene sentido contarla
        if not plaza:
            continue

        if forma_raw in ("mensual", "m"):
            plazas_mensual.add(plaza)
        elif forma_raw in ("trimestral", "t"):
            plazas_trimestral.add(plaza)
        elif forma_raw in ("semestral", "s"):
            plazas_semestral.add(plaza)
        elif forma_raw in ("anual", "a"):
            plazas_anual.add(plaza)
        else:
            plazas_otro.add(plaza)

    total_pago_mensual = len(plazas_mensual)
    total_pago_trimestral = len(plazas_trimestral)
    total_pago_semestral = len(plazas_semestral)
    total_pago_anual = len(plazas_anual)
    total_pago_otro = len(plazas_otro)

    # (Opcional) Log de depuración
    current_app.logger.info(
        "Contadores forma_pago: mensual=%s, trimestral=%s, "
        "semestral=%s, anual=%s, otros=%s",
        total_pago_mensual,
        total_pago_trimestral,
        total_pago_semestral,
        total_pago_anual,
        total_pago_otro,
    )

    # 5) Construcción del documento Word
    document = Document()
    section = document.sections[0]
    hoy = date.today()

    # 5.1 Header con resumen de plazas (global)
    total_plazas_mensual_global = sum(
        a.get("total_plazas", 0) or 0 for a in padron_principal
    )

    resumen_plazas = _obtener_resumen_plazas_rio_torio(
        fecha_inicio_periodo,
        fecha_fin_periodo,
    )

    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_paragraph.add_run("Jefatura de Policía Local\nAdministración\n")
    header_paragraph.add_run(
        f"Plazas totales: {total_plazas_mensual_global} | "
        f"Libres: {resumen_plazas['libres']} | "
        f"Ocupadas: {resumen_plazas['ocupadas']}"
    )

    section.left_margin = Cm(3)

    # 5.2 Asunto
    p_asunto = document.add_paragraph()
    p_asunto.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_asunto.add_run(
        "ASUNTO: Relación de abonados y plazas del parking de camiones de la "
        "Calle Río Torío (formato conjunto)."
    ).bold = True

    document.add_paragraph()

    # 5.3 Interesado
    p_interesado = document.add_paragraph()
    p_interesado.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_interesado.add_run("Interesado: Gestión Tributaria.")

    document.add_paragraph()

    # 5.4 Párrafo de comunicación
    p_comunico = document.add_paragraph()
    p_comunico.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_comunico.paragraph_format.first_line_indent = Cm(2)
    p_comunico.add_run(
        f"Comunico a Vd. que durante el mes de {nombre_mes} de {anio_informe} "
        "la relación de abonados y sus plazas es la siguiente:"
    )
    document.add_paragraph()

    # 5.5 Contadores en el inicio (ANTES de la tabla)
    p_contadores = document.add_paragraph()
    p_contadores.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_contadores.paragraph_format.first_line_indent = Cm(2)
    texto_contadores = (
        f"total pago mensual = {total_pago_mensual} | "
        f"trimestral = {total_pago_trimestral} | "
        f"semestral = {total_pago_semestral} | "
        f"anual = {total_pago_anual}"
    )
    if total_pago_otro:
        texto_contadores += f" | otros = {total_pago_otro}"
    p_contadores.add_run(texto_contadores).bold = True

    document.add_paragraph()

    # 5.6 Tabla principal (UNA SOLA TABLA, 5 COLUMNAS)
    if padron_principal:
        table = document.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Apellidos y nombre"
        hdr_cells[1].text = "NIF"
        hdr_cells[2].text = "Forma de pago"
        hdr_cells[3].text = "Plazas"
        hdr_cells[4].text = "Total plazas"

        for a in padron_principal:
            row = table.add_row().cells
            row[0].text = str(a.get("nombre_proveedor", ""))
            row[1].text = str(a.get("nif", ""))
            row[2].text = (a.get("forma_pago") or "").upper()
            row[3].text = ", ".join(a.get("plazas", []))
            row[4].text = str(a.get("total_plazas", 0) or 0)
    else:
        p_sin = document.add_paragraph()
        p_sin.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_sin.add_run("Actualmente no existen abonados registrados.").italic = True

    document.add_paragraph()

    # 5.7 Variaciones (solo si hay)
    if variaciones:
        p_var_titulo = document.add_paragraph()
        p_var_titulo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_var_titulo.paragraph_format.first_line_indent = Cm(2)
        p_var_titulo.add_run(
            "Variaciones (bajas, cambios de plaza, cambios de forma de pago):"
        ).bold = True

        table_var = document.add_table(rows=1, cols=5)
        table_var.style = "Table Grid"
        hdrv = table_var.rows[0].cells
        hdrv[0].text = "Fecha"
        hdrv[1].text = "Proveedor"
        hdrv[2].text = "NIF"
        hdrv[3].text = "Forma de pago"
        hdrv[4].text = "Detalle"

        for v in variaciones:
            rowv = table_var.add_row().cells
            fv = v.get("fecha")
            rowv[0].text = fv.strftime("%d/%m/%Y") if fv else ""
            rowv[1].text = v.get("nombre_proveedor", "")
            rowv[2].text = v.get("nif", "")
            rowv[3].text = (v.get("forma_pago") or "").upper()
            rowv[4].text = f"{v.get('tipo_cambio', '')}: {v.get('numero_plaza', '')}"
    else:
        p_sin_var = document.add_paragraph()
        p_sin_var.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_sin_var.paragraph_format.first_line_indent = Cm(2)
        p_sin_var.add_run("SIN VARIACIONES EN ESTE MES.").bold = True

    document.add_paragraph()

    # 5.8 Cierre
    p_final = document.add_paragraph()
    p_final.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_final.paragraph_format.first_line_indent = Cm(2)
    p_final.add_run(
        "Lo que comunico a Vd., para su conocimiento y efectos oportunos."
    )

    document.add_paragraph()
    p_firma_fecha = document.add_paragraph()
    p_firma_fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_firma_fecha.add_run(f"En Ávila, a {hoy.strftime('%d/%m/%Y')}.")

    document.add_paragraph()
    p_cargo = document.add_paragraph()
    p_cargo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cargo.add_run("El Jefe de la Policía Local.")

    document.add_paragraph()
    p_nombre = document.add_paragraph()
    p_nombre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_nombre.add_run("D. Carlos Blanco Rubio")

    p_digital = document.add_paragraph()
    p_digital.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_digital.add_run("(documento firmado digitalmente)").italic = True

    # 6) Guardar y devolver ruta
    document.save(ruta_completa)
    return ruta_completa


# =============================================================================
# 9️⃣ ENDPOINT JSON: PADRÓN TÉCNICO
# =============================================================================

@btn_rio_torio_padron_manual_bp.route("/padron/generar", methods=["GET"])
@login_required
@rol_required("gestor", "super_admin")
def btn_rio_torio_padron_generar_json():
    """
    Endpoint técnico JSON para comprobar los datos del padrón.

    - Calcula el periodo del padrón.
    - Devuelve:
        · periodo (inicio, fin, año, mes, nombre).
        · padron_principal.
        · variaciones.
    """
    periodo = obtener_periodo_padron()
    inicio_mes = periodo["inicio_mes"]
    fin_mes = periodo["fin_mes"]

    current_app.logger.info(
        "Generando padrón Río Torío (JSON) para periodo %s - %s",
        inicio_mes,
        fin_mes,
    )

    filas = obtener_historico_periodo(inicio_mes, fin_mes)
    padron_principal = construir_padron_principal(filas)
    variaciones = construir_variaciones(filas, inicio_mes, fin_mes)

    return {
        "periodo": {
            "inicio": inicio_mes.isoformat(),
            "fin": fin_mes.isoformat(),
            "anio": periodo["anio"],
            "mes": periodo["mes_num"],
            "mes_nombre": periodo["mes_nombre"],
        },
        "padron_principal": padron_principal,
        "variaciones": variaciones,
    }


# =============================================================================
# 🔟 BOTÓN CON FORMULARIO MES/AÑO → WORD
#   Acceso restringido a roles: gestor, super_admin
# =============================================================================

@btn_rio_torio_padron_manual_bp.route(
    "/btn_rio_torio_padron_manual",
    methods=["GET", "POST"],
)
@login_required
@rol_required("gestor", "super_admin")
def btn_rio_torio_padron_manual():
    """
    Botón de super admin para generar el padrón manual (mes/año).

    - GET:
        · Calcula por defecto el mes anterior al actual.
        · Muestra formulario con select de MES/AÑO.
    - POST:
        · Valida MES/AÑO.
        · Genera el Word.
        · Devuelve el fichero como descarga.
    """
    current_app.logger.info(
        "Entrando en btn_rio_torio_padron_manual (padron manual)."
    )

    # GET: mostrar formulario con mes/año por defecto
    if request.method == "GET":
        hoy = date.today()
        if hoy.month == 1:
            mes_defecto = 12
            anio_defecto = hoy.year - 1
        else:
            mes_defecto = hoy.month - 1
            anio_defecto = hoy.year

        return render_template(
            "parquin/rio_torio/rio_torio_padron_manual.html",
            meses=MESES_ES,
            mes_defecto=mes_defecto,
            anio_defecto=anio_defecto,
        )

    # POST: procesar formulario
    mes_str = request.form.get("mes")
    anio_str = request.form.get("anio")

    # Validar MES/AÑO
    try:
        mes = int(mes_str)
        anio = int(anio_str)
        if mes < 1 or mes > 12:
            raise ValueError("Mes fuera de rango")
    except Exception:
        flash("Mes o año no válidos.", "danger")
        return redirect(
            url_for(
                "btn_rio_torio_padron_manual_bp.btn_rio_torio_padron_manual"
            )
        )

    # Generar informe Word
    try:
        ruta = _generar_informe_word(mes, anio)
        flash(f"Padrón generado correctamente: {ruta.name}", "success")

        return send_file(
            ruta,
            as_attachment=True,
            download_name=ruta.name,
            mimetype=(
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            ),
        )
    except Exception as e:
        current_app.logger.error(f"Error generando padrón Rio Torío: {e}")
        flash("Error generando el informe de padrón.", "danger")
        return redirect(
            url_for(
                "btn_rio_torio_padron_manual_bp.btn_rio_torio_padron_manual"
            )
        )
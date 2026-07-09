# =============================================================================
# MÓDULO: PADRÓN RÍO TORÍO
# -----------------------------------------------------------------------------
# Autor: thierro23-cloud
# Fecha: 2026-07-09
# Versión: 3.0 (Documento Word profesional + resumen operativo)
#
# DESCRIPCIÓN
# Este módulo:
#   1) Calcula el periodo del padrón.
#   2) Recupera histórico activo del periodo para Río Torío.
#   3) Construye:
#      - Tabla principal (abonados por proveedor).
#      - Variaciones (baja, cambio de plaza ±7 días, forma de pago).
#   4) Genera automáticamente documento Word desde plantilla "normal":
#      - Encabezado derecho institucional.
#      - Contador de plazas (totales, libres, ocupadas).
#      - Asunto/Destinatario y texto justificado.
#      - Contadores de usuarios mensual/anual (ANUAL en rojo).
#      - Tabla: Apellidos | Nombre | NIF | Forma de pago.
#      - Cierre de firma electrónica.
#
# NOTAS DE IMPLANTACIÓN
# - Requiere python-docx.
# - Ajustar rutas de plantilla/salida según despliegue.
# - Si en plantilla existen placeholders {{...}}, se reemplazan automáticamente.
# =============================================================================

from __future__ import annotations

from datetime import date, datetime, timedelta
from collections import defaultdict
from pathlib import Path
import calendar

from flask import Blueprint, current_app, jsonify
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor, Pt

from db import ejecutar_query


# =============================================================================
# CONFIGURACIÓN
# =============================================================================
VENTANA_CAMBIO_PLAZA_DIAS = 7

# Plantilla "normal" (ajusta nombre/ruta real si difiere)
PLANTILLA_WORD = Path("templates/parquin/plantilla_padron_normal.docx")

# Carpeta de salida para documentos
SALIDA_PADRON = Path("generated/parquin/padron")

# Mapeo mes español
MESES_ES = {
    1: "enero",
    2: "febrero",
    3: "marzo",
    4: "abril",
    5: "mayo",
    6: "junio",
    7: "julio",
    8: "agosto",
    9: "septiembre",
    10: "octubre",
    11: "noviembre",
    12: "diciembre",
}


# =============================================================================
# BLUEPRINT
# =============================================================================
rio_torio_padron_bp = Blueprint(
    "rio_torio_padron_bp",
    __name__,
    url_prefix="/parquin/rio_torio/padron",
)


# =============================================================================
# 1) CÁLCULO DEL PERIODO
# =============================================================================
def obtener_periodo_padron(fecha_base: date | None = None) -> dict:
    """
    Calcula el periodo del padrón según la fecha base.
    """
    if fecha_base is None:
        fecha_base = date.today()

    if fecha_base.month == 12:
        anio_siguiente = fecha_base.year + 1
        mes_siguiente = 1
    else:
        anio_siguiente = fecha_base.year
        mes_siguiente = fecha_base.month + 1

    primer_dia_mes_siguiente = date(anio_siguiente, mes_siguiente, 1)

    d = primer_dia_mes_siguiente
    while d.weekday() >= 5:  # 5=sábado, 6=domingo
        d += timedelta(days=1)
    primer_laborable = d

    # Mes del padrón = mes anterior al mes_siguiente
    if mes_siguiente == 1:
        anio_padron = anio_siguiente - 1
        mes_padron = 12
    else:
        anio_padron = anio_siguiente
        mes_padron = mes_siguiente - 1

    inicio_mes = date(anio_padron, mes_padron, 1)
    _, ultimo_dia = calendar.monthrange(anio_padron, mes_padron)
    fin_mes = date(anio_padron, mes_padron, ultimo_dia)

    return {
        "inicio_mes": inicio_mes,
        "fin_mes": fin_mes,
        "anio": anio_padron,
        "mes_num": mes_padron,
        "mes_nombre_en": calendar.month_name[mes_padron],
        "mes_nombre_es": MESES_ES.get(mes_padron, str(mes_padron)),
        "primer_laborable_mes_siguiente": primer_laborable,
    }


# =============================================================================
# 2) CONSULTAS
# =============================================================================
def obtener_historico_periodo(inicio_mes: date, fin_mes: date) -> list[dict]:
    """
    Recupera histórico activo en [inicio_mes, fin_mes] con proveedor/plaza/pago.
    """
    query = """
        SELECT
            h.idtbl_historico_plazas,
            h.idtbl_plazas,
            h.fecha_inicio,
            h.fecha_fin,
            h.exp_solicitud_fin,
            h.forma_pago,
            pl.codigo_plazas,

            h.idtbl_proveedores AS idtbl_usuarios,
            h.idtbl_proveedores,
            h.forma_pago AS forma_pago_usuario,

            COALESCE(
                NULLIF(TRIM(CONCAT_WS(' ', p.apellidos, p.nombre)), ''),
                NULLIF(TRIM(p.nombre_razon_social), ''),
                'PROVEEDOR SIN NOMBRE'
            ) AS nombre_proveedor,

            p.NIF AS nif_proveedor,
            p.apellidos AS apellidos_proveedor,
            p.nombre AS nombre_proveedor_individual,
            p.nombre_razon_social AS nombre_razon_social_proveedor
        FROM parquin_camiones.tbl_historico_plazas AS h
        JOIN parquin_camiones.tbl_plazas AS pl
          ON h.idtbl_plazas = pl.idtbl_plazas
        JOIN bd_tbl_comunes.tbl_proveedores AS p
          ON h.idtbl_proveedores = p.Idtbl_proveedores
        WHERE
          h.fecha_inicio <= %s
          AND (h.fecha_fin IS NULL OR h.fecha_fin >= %s)
        ORDER BY nombre_proveedor, pl.codigo_plazas, h.fecha_inicio
    """

    return ejecutar_query(
        query,
        params=(fin_mes, inicio_mes),
        nombre_bd="parquin_camiones",
    )


def obtener_total_plazas_catalogo() -> int:
    """
    Total de plazas físicas del catálogo.
    """
    query = """
        SELECT COUNT(*) AS total_plazas
        FROM parquin_camiones.tbl_plazas
    """
    filas = ejecutar_query(query, params=(), nombre_bd="parquin_camiones")
    if not filas:
        return 0
    return int(filas[0].get("total_plazas") or 0)


# =============================================================================
# 3) CONSTRUCCIÓN DE TABLAS LÓGICAS
# =============================================================================
def normalizar_forma_pago(valor: str | None) -> str:
    """
    Normaliza forma de pago a: 'MENSUAL', 'ANUAL' u 'OTRO'.
    """
    if not valor:
        return "OTRO"

    v = str(valor).strip().lower()
    if "anual" in v or "año" in v or "anio" in v:
        return "ANUAL"
    if "mensual" in v or "mes" == v:
        return "MENSUAL"
    return "OTRO"


def construir_padron_principal(filas: list[dict]) -> list[dict]:
    """
    Agrupa por proveedor y consolida datos principales.
    """
    por_proveedor: dict[int, dict] = {}

    for f in filas:
        id_prov = f["idtbl_proveedores"]
        forma_pago = (f.get("forma_pago") or f.get("forma_pago_usuario") or "").strip()

        if id_prov not in por_proveedor:
            por_proveedor[id_prov] = {
                "idtbl_proveedores": id_prov,
                "nombre_proveedor": (f.get("nombre_proveedor") or "").strip(),
                "apellidos": (f.get("apellidos_proveedor") or "").strip(),
                "nombre": (f.get("nombre_proveedor_individual") or "").strip(),
                "razon_social": (f.get("nombre_razon_social_proveedor") or "").strip(),
                "nif": (f.get("nif_proveedor") or "").strip(),
                "forma_pago": forma_pago,
                "plazas": set(),
            }

        por_proveedor[id_prov]["plazas"].add(f.get("codigo_plazas"))

        # Si aún no tiene forma de pago, capturamos la primera no vacía
        if not por_proveedor[id_prov]["forma_pago"] and forma_pago:
            por_proveedor[id_prov]["forma_pago"] = forma_pago

    resultado = []
    for prov in por_proveedor.values():
        plazas_ordenadas = sorted([p for p in prov["plazas"] if p is not None], key=str)
        prov["plazas"] = plazas_ordenadas
        prov["total_plazas"] = len(plazas_ordenadas)
        prov["forma_pago_normalizada"] = normalizar_forma_pago(prov.get("forma_pago"))
        resultado.append(prov)

    resultado.sort(
        key=lambda x: (
            (x.get("apellidos") or "").lower(),
            (x.get("nombre") or "").lower(),
            (x.get("razon_social") or "").lower(),
            (x.get("nif") or "").lower(),
        )
    )
    return resultado


def _score_emparejamiento_cambio(fecha_baja: date, fecha_alta: date, ventana_dias: int):
    """
    Score:
      (0, d) para alta posterior/igual dentro de ventana
      (1, d) para alta anterior dentro de ventana
      None fuera de ventana
    """
    delta = (fecha_alta - fecha_baja).days
    if 0 <= delta <= ventana_dias:
        return (0, delta)
    if -ventana_dias <= delta < 0:
        return (1, abs(delta))
    return None


def construir_variaciones(
    filas: list[dict],
    inicio_mes: date,
    fin_mes: date,
    ventana_cambio_dias: int = VENTANA_CAMBIO_PLAZA_DIAS,
) -> list[dict]:
    """
    Variaciones:
      - BAJA
      - Cambio de plaza (emparejamiento contrariado/complementario en ±ventana)
      - Forma de pago
    """
    variaciones: list[dict] = []
    por_prov_usuario: dict[tuple[int, int], list[dict]] = defaultdict(list)

    for f in filas:
        clave = (f["idtbl_proveedores"], f["idtbl_usuarios"])
        por_prov_usuario[clave].append(f)

    for (_id_prov, _id_usuario), registros in por_prov_usuario.items():
        registros.sort(
            key=lambda r: (
                r["fecha_inicio"] or date.min,
                r["fecha_fin"] or date.max,
                r["idtbl_plazas"],
            )
        )

        bajas: list[dict] = []
        altas: list[dict] = []

        for r in registros:
            fi = r["fecha_inicio"]
            ff = r["fecha_fin"]

            if fi is not None and inicio_mes <= fi <= fin_mes:
                altas.append(r)

            if ff is not None and inicio_mes <= ff <= fin_mes:
                bajas.append(r)

        usadas_altas_idx: set[int] = set()
        usadas_bajas_id: set[int] = set()

        # Cambio de plaza por variación contraria/complementaria
        for baja in bajas:
            ff = baja["fecha_fin"]
            if ff is None:
                continue

            mejor_idx = None
            mejor_score = None

            for idx_alta, alta in enumerate(altas):
                if idx_alta in usadas_altas_idx:
                    continue
                if alta["idtbl_plazas"] == baja["idtbl_plazas"]:
                    continue

                fi = alta["fecha_inicio"]
                if fi is None:
                    continue

                score = _score_emparejamiento_cambio(ff, fi, ventana_cambio_dias)
                if score is None:
                    continue

                if mejor_score is None or score < mejor_score:
                    mejor_score = score
                    mejor_idx = idx_alta

            if mejor_idx is not None:
                alta = altas[mejor_idx]
                variaciones.append(
                    {
                        "nombre_proveedor": baja["nombre_proveedor"],
                        "nif": baja["nif_proveedor"],
                        "codigo_plazas": f"{baja['codigo_plazas']} -> {alta['codigo_plazas']}",
                        "tipo_cambio": "Cambio de plaza",
                        "fecha": ff,
                    }
                )
                usadas_bajas_id.add(baja["idtbl_historico_plazas"])
                usadas_altas_idx.add(mejor_idx)

        # Bajas puras
        for baja in bajas:
            if baja["idtbl_historico_plazas"] in usadas_bajas_id:
                continue
            variaciones.append(
                {
                    "nombre_proveedor": baja["nombre_proveedor"],
                    "nif": baja["nif_proveedor"],
                    "codigo_plazas": baja["codigo_plazas"],
                    "tipo_cambio": "BAJA",
                    "fecha": baja["fecha_fin"],
                }
            )

        # Cambio forma de pago
        por_plaza: dict[int, list[dict]] = defaultdict(list)
        for r in registros:
            por_plaza[r["idtbl_plazas"]].append(r)

        for _id_plaza, regs_plaza in por_plaza.items():
            regs_plaza.sort(key=lambda r: r["fecha_inicio"] or date.min)
            forma_anterior = None

            for r in regs_plaza:
                forma_actual = (r.get("forma_pago") or r.get("forma_pago_usuario") or "").strip()
                fi = r["fecha_inicio"]

                if forma_anterior is not None and forma_actual != forma_anterior:
                    if fi is not None and inicio_mes <= fi <= fin_mes:
                        variaciones.append(
                            {
                                "nombre_proveedor": r["nombre_proveedor"],
                                "nif": r["nif_proveedor"],
                                "codigo_plazas": r["codigo_plazas"],
                                "tipo_cambio": "Forma de pago",
                                "fecha": fi,
                            }
                        )

                forma_anterior = forma_actual

    variaciones.sort(key=lambda v: (v["nombre_proveedor"], v["fecha"], v["codigo_plazas"]))
    return variaciones


# =============================================================================
# 4) WORD: UTILIDADES
# =============================================================================
def _replace_text_everywhere(doc: Document, replacements: dict[str, str]) -> None:
    """
    Reemplaza placeholders {{CLAVE}} en párrafos, celdas de tablas y cabecera.
    """
    def _replace_in_paragraph(paragraph):
        if not paragraph.runs:
            txt = paragraph.text
            for k, v in replacements.items():
                txt = txt.replace(f"{{{{{k}}}}}", str(v))
            paragraph.text = txt
            return

        full_text = "".join(run.text for run in paragraph.runs)
        new_text = full_text
        for k, v in replacements.items():
            new_text = new_text.replace(f"{{{{{k}}}}}", str(v))
        if new_text != full_text:
            for run in paragraph.runs:
                run.text = ""
            paragraph.runs[0].text = new_text

    # body
    for p in doc.paragraphs:
        _replace_in_paragraph(p)

    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph(p)

    # headers/footers
    for section in doc.sections:
        for p in section.header.paragraphs:
            _replace_in_paragraph(p)
        for t in section.header.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        _replace_in_paragraph(p)

        for p in section.footer.paragraphs:
            _replace_in_paragraph(p)


def _set_header_institucional(
    doc: Document,
    total_plazas: int,
    libres: int,
    ocupadas: int,
) -> None:
    """
    Encabezado a la derecha:
      Jefatura de Policía Local
      Administración
      Plazas totales: X
      Plazas libres: Y
      Plazas ocupadas: Z
    """
    for section in doc.sections:
        header = section.header

        # Limpia párrafos previos para evitar duplicados en regeneraciones
        for p in list(header.paragraphs):
            p._element.getparent().remove(p._element)

        lineas = [
            "Jefatura de Policía Local",
            "Administración",
            f"Plazas totales: {total_plazas}",
            f"Plazas libres: {libres}",
            f"Plazas ocupadas: {ocupadas}",
        ]

        for linea in lineas:
            p = header.add_paragraph(linea)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _agregar_bloque_asunto_destinatario(doc: Document, mes_texto: str) -> None:
    """
    Añade bloque justificado de asunto/destinatario/comunicación.
    """
    p1 = doc.add_paragraph(
        "Asunto: Comunicación de las modificaciones de los abonados del parking de camiones calle Río Torio"
    )
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p2 = doc.add_paragraph("Destinatario: Gestión tributaria")
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p3 = doc.add_paragraph(
        f"Comunico a usted que durante el mes de {mes_texto}, la relación de abonados es la siguiente:"
    )
    p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def _agregar_resumen_pago(doc: Document, n_mensual: int, n_anual: int) -> None:
    """
    Añade resumen encima de la tabla:
      Mensual: X
      ANUAL: Y   (ANUAL en rojo)
    """
    p_m = doc.add_paragraph()
    p_m.add_run(f"Mensual: {n_mensual}")

    p_a = doc.add_paragraph()
    run_lbl = p_a.add_run("ANUAL: ")
    run_lbl.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # rojo
    run_val = p_a.add_run(str(n_anual))
    run_val.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # rojo


def _agregar_tabla_abonados(doc: Document, padron_principal: list[dict]) -> None:
    """
    Tabla: Apellidos | Nombre | NIF | Forma de pago
    """
    tabla = doc.add_table(rows=1, cols=4)
    tabla.style = "Table Grid"

    hdr = tabla.rows[0].cells
    hdr[0].text = "Apellidos"
    hdr[1].text = "Nombre"
    hdr[2].text = "NIF"
    hdr[3].text = "Forma de pago"

    # Orden profesional por apellidos/nombre/razón social
    datos = sorted(
        padron_principal,
        key=lambda x: (
            (x.get("apellidos") or "").lower(),
            (x.get("nombre") or "").lower(),
            (x.get("razon_social") or "").lower(),
            (x.get("nif") or "").lower(),
        ),
    )

    for p in datos:
        apellidos = (p.get("apellidos") or "").strip()
        nombre = (p.get("nombre") or "").strip()
        razon_social = (p.get("razon_social") or "").strip()

        # Si no hay persona física, usar razón social en columna Nombre
        if not apellidos and not nombre and razon_social:
            nombre = razon_social

        nif = (p.get("nif") or "").strip()
        fp = (p.get("forma_pago_normalizada") or p.get("forma_pago") or "").strip()

        row = tabla.add_row().cells
        row[0].text = apellidos
        row[1].text = nombre
        row[2].text = nif
        row[3].text = fp


def _agregar_bloque_firma(doc: Document, fecha_firma: date) -> None:
    """
    Cierre:
      Ávila a <fecha firma>
      firma electrónica
      El Jefe de la Policía Local
      D. Carlos Blanco Rubio
      documento firmado electrónicamente (pequeño)
    """
    doc.add_paragraph("")  # espacio
    doc.add_paragraph(f"Ávila a {fecha_firma.strftime('%d/%m/%Y')}")

    doc.add_paragraph("firma electrónica")
    doc.add_paragraph("El Jefe de la Policía Local")
    doc.add_paragraph("D. Carlos Blanco Rubio")

    p_small = doc.add_paragraph("documento firmado electrónicamente")
    if p_small.runs:
        p_small.runs[0].font.size = Pt(8)


def generar_word_padron(
    periodo: dict,
    padron_principal: list[dict],
    total_plazas: int,
    fecha_firma: date | None = None,
) -> Path:
    """
    Genera DOCX final profesional cumpliendo estructura solicitada.
    """
    if fecha_firma is None:
        fecha_firma = date.today()

    if not PLANTILLA_WORD.exists():
        raise FileNotFoundError(f"No existe la plantilla Word: {PLANTILLA_WORD}")

    SALIDA_PADRON.mkdir(parents=True, exist_ok=True)

    nombre_archivo = (
        f"{periodo['anio']}_padron_{periodo['mes_num']:02d}_rio_torio_"
        f"{fecha_firma.strftime('%Y%m%d')}.docx"
    )
    ruta_salida = SALIDA_PADRON / nombre_archivo

    doc = Document(str(PLANTILLA_WORD))

    # Cálculos operativos
    ocupadas = sum(1 for p in padron_principal if (p.get("total_plazas") or 0) > 0)
    libres = max(total_plazas - ocupadas, 0)

    n_mensual = sum(
        1 for p in padron_principal if (p.get("forma_pago_normalizada") == "MENSUAL")
    )
    n_anual = sum(
        1 for p in padron_principal if (p.get("forma_pago_normalizada") == "ANUAL")
    )

    # Reemplazo de placeholders si plantilla los incluye
    _replace_text_everywhere(
        doc,
        {
            "ANIO": str(periodo["anio"]),
            "MES_NUM": f"{periodo['mes_num']:02d}",
            "MES_NOMBRE_ES": str(periodo["mes_nombre_es"]),
            "FECHA_INICIO": periodo["inicio_mes"].isoformat(),
            "FECHA_FIN": periodo["fin_mes"].isoformat(),
            "PLAZAS_TOTALES": str(total_plazas),
            "PLAZAS_LIBRES": str(libres),
            "PLAZAS_OCUPADAS": str(ocupadas),
            "N_MENSUAL": str(n_mensual),
            "N_ANUAL": str(n_anual),
            "FECHA_FIRMA": fecha_firma.strftime("%d/%m/%Y"),
        },
    )

    # Encabezado institucional derecho
    _set_header_institucional(
        doc=doc,
        total_plazas=total_plazas,
        libres=libres,
        ocupadas=ocupadas,
    )

    # Cuerpo principal
    mes_texto = f"{periodo['mes_nombre_es']} de {periodo['anio']}"
    _agregar_bloque_asunto_destinatario(doc, mes_texto)
    _agregar_resumen_pago(doc, n_mensual=n_mensual, n_anual=n_anual)
    _agregar_tabla_abonados(doc, padron_principal)
    _agregar_bloque_firma(doc, fecha_firma=fecha_firma)

    doc.save(str(ruta_salida))
    return ruta_salida


# =============================================================================
# 5) ENDPOINT PRINCIPAL
# =============================================================================
@rio_torio_padron_bp.route("/generar", methods=["GET"])
def generar_padron_rio_torio():
    """
    Genera padrón + Word profesional.
    La fecha de cierre/firma es la fecha actual de ejecución.
    """
    periodo = obtener_periodo_padron()
    inicio_mes = periodo["inicio_mes"]
    fin_mes = periodo["fin_mes"]

    current_app.logger.info(
        "Generando padrón Río Torío periodo %s - %s (ventana cambio plaza ±%s días)",
        inicio_mes,
        fin_mes,
        VENTANA_CAMBIO_PLAZA_DIAS,
    )

    filas = obtener_historico_periodo(inicio_mes, fin_mes)
    padron_principal = construir_padron_principal(filas)
    variaciones = construir_variaciones(
        filas=filas,
        inicio_mes=inicio_mes,
        fin_mes=fin_mes,
        ventana_cambio_dias=VENTANA_CAMBIO_PLAZA_DIAS,
    )

    total_plazas = obtener_total_plazas_catalogo()
    fecha_firma = date.today()  # fecha de firma = fecha de ejecución

    ruta_word = generar_word_padron(
        periodo=periodo,
        padron_principal=padron_principal,
        total_plazas=total_plazas,
        fecha_firma=fecha_firma,
    )

    return jsonify(
        {
            "ok": True,
            "periodo": {
                "inicio": inicio_mes.isoformat(),
                "fin": fin_mes.isoformat(),
                "anio": periodo["anio"],
                "mes": periodo["mes_num"],
                "mes_nombre_es": periodo["mes_nombre_es"],
                "primer_laborable_mes_siguiente": periodo["primer_laborable_mes_siguiente"].isoformat(),
            },
            "firma": {
                "fecha_firma": fecha_firma.isoformat(),
            },
            "configuracion": {
                "ventana_cambio_plaza_dias": VENTANA_CAMBIO_PLAZA_DIAS,
                "plantilla_word": str(PLANTILLA_WORD),
                "salida_word_dir": str(SALIDA_PADRON),
            },
            "resumen": {
                "total_plazas_catalogo": total_plazas,
                "ocupadas": sum(1 for p in padron_principal if (p.get("total_plazas") or 0) > 0),
                "libres": max(total_plazas - sum(1 for p in padron_principal if (p.get("total_plazas") or 0) > 0), 0),
                "usuarios_mensual": sum(1 for p in padron_principal if p.get("forma_pago_normalizada") == "MENSUAL"),
                "usuarios_anual": sum(1 for p in padron_principal if p.get("forma_pago_normalizada") == "ANUAL"),
            },
            "documento_word": str(ruta_word),
            "padron_principal": padron_principal,
            "variaciones": variaciones,
        }
    )

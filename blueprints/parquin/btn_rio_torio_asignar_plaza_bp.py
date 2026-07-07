from __future__ import annotations

"""
===============================================================================
MÓDULO: BTN_RIO_TORIO_ASIGNAR_PLAZA_BP
===============================================================================

OBJETIVO GENERAL
----------------
Este módulo gestiona la asignación y liberación de plazas del parking Río Torío
y, como parte del mismo flujo funcional, genera el informe de adjudicación en
formato Word para dejar constancia documental de cada operación.

FLUJO DE NEGOCIO
----------------
1. El usuario gestor selecciona una o varias plazas.
2. Puede asignarlas a un usuario/proveedor o liberarlas.
3. En el mismo acto se guarda el histórico en tbl_historico_plazas.
4. Si se pulsa el botón de informe:
   - se genera el documento .docx;
   - se crea una copia .pdf en blueprints/parquin/padron;
   - se devuelve el .docx para descarga/visualización.

REGLAS CLAVE
------------
- Si no se indica fecha_asignacion, se usa datetime.datetime.now().
- Si no se indica fecha_liberacion, se usa datetime.datetime.now().
- Al liberar una plaza, la fecha indicada se guarda como fecha_fin en
  tbl_historico_plazas.
- Al asignar una plaza, la fecha indicada se guarda como fecha_inicio.
- El informe debe poder generarse siempre que se pulse el botón, tanto si se
  asigna como si se libera.
- Se mantiene un documento Word como evidencia y una copia PDF archivada.

SEGURIDAD
---------
- login_required
- rol_required("gestor", "super_admin")

DEPENDENCIAS
------------
- ejecutar_query, ejecutar_non_query, get_connection
- python-docx
- conversión a PDF mediante LibreOffice en modo headless si está disponible
===============================================================================
"""

from flask import (
    Blueprint,
    render_template,
    request,
    redirect,
    url_for,
    flash,
    current_app,
    send_file,
)
from services.helpers import login_required, rol_required
from db import ejecutar_query, ejecutar_non_query, get_connection

import datetime
import subprocess
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

btn_rio_torio_asignar_plaza_bp = Blueprint(
    "btn_rio_torio_asignar_plaza_bp",
    __name__,
    url_prefix="/parquin/rio_torio",
)


# =============================================================================
# CONFIGURACIÓN DE RUTAS Y ARCHIVOS
# =============================================================================

BASE_DIR = Path(__file__).resolve().parent.parent.parent
PADRON_DIR = BASE_DIR / "parquin" / "padron"
PADRON_DIR.mkdir(parents=True, exist_ok=True)


# =============================================================================
# FUNCIONES AUXILIARES
# =============================================================================


def _get_orden_plazas():
    """
    Determina el criterio de ordenación solicitado por querystring.

    Devuelve:
        tuple[str, str, str, str]:
            - columna SQL segura
            - dirección ASC/DESC
            - nombre lógico de sort
            - valor original dir
    """
    columnas_validas = {
        "fila": "p.fila",
        "codigo": "p.codigo_plazas",
        "proveedor": "proveedor_actual",
        "cuenta": "u.numero_cuenta",
    }

    sort = request.args.get("sort", "fila")
    dir_param = request.args.get("dir", "asc").lower()
    columna_order = columnas_validas.get(sort, "p.fila")
    direction = "DESC" if dir_param == "desc" else "ASC"
    return columna_order, direction, sort, dir_param


def _cargar_datos_parquin_rio_torio():
    """
    Carga usuarios activos y plazas para la pantalla principal.

    Devuelve:
        usuarios_parquin, plazas, sort, dir_param
    """
    usuarios_parquin = ejecutar_query(
        """
        SELECT
            u.idtbl_usuarios,
            u.idtbl_proveedores,
            u.numero_cuenta,
            u.activo_baja,
            u.rol,
            p.Nombre_Razon_Social AS nombre_proveedor
        FROM tbl_usuarios AS u
        INNER JOIN bd_tbl_comunes.tbl_proveedores AS p
            ON u.idtbl_proveedores = p.Idtbl_proveedores
        WHERE u.activo_baja = 1
        ORDER BY p.Nombre_Razon_Social
        """,
        nombre_bd="parquin_camiones",
    )

    columna_order, direction, sort, dir_param = _get_orden_plazas()

    plazas = ejecutar_query(
        f"""
        SELECT
            p.idtbl_plazas,
            p.codigo_plazas,
            p.fila,
            u.idtbl_usuarios,
            pr.Nombre_Razon_Social AS proveedor_actual,
            u.numero_cuenta,
            pr.Persona_contacto_comercial AS persona_contacto_comercial,
            pr.Persona_contacto_admin AS persona_contacto_admin
        FROM tbl_plazas AS p
        LEFT JOIN tbl_usuarios AS u
            ON p.idtbl_usuarios = u.idtbl_usuarios
        LEFT JOIN bd_tbl_comunes.tbl_proveedores AS pr
            ON u.idtbl_proveedores = pr.Idtbl_proveedores
        ORDER BY {columna_order} {direction}, p.idtbl_plazas
        """,
        nombre_bd="parquin_camiones",
    )

    return usuarios_parquin, plazas, sort, dir_param


def _mes_siguiente(fecha: datetime.date) -> datetime.date:
    """
    Devuelve el día 1 del mes siguiente a la fecha indicada.
    """
    if fecha.month == 12:
        return datetime.date(fecha.year + 1, 1, 1)
    return datetime.date(fecha.year, fecha.month + 1, 1)


def _generar_pdf_desde_docx(ruta_docx: Path, ruta_pdf: Path) -> bool:
    """
    Intenta convertir un DOCX a PDF usando LibreOffice en modo headless.

    Returns:
        bool: True si la conversión parece exitosa, False en caso contrario.
    """
    comando = [
        "soffice",
        "--headless",
        "--convert-to",
        "pdf",
        "--outdir",
        str(ruta_pdf.parent),
        str(ruta_docx),
    ]

    try:
        resultado = subprocess.run(
            comando,
            capture_output=True,
            text=True,
            check=False,
        )
        pdf_generado = ruta_docx.with_suffix(".pdf")
        if pdf_generado.exists():
            if pdf_generado != ruta_pdf:
                pdf_generado.replace(ruta_pdf)
            return True
        current_app.logger.warning(
            "Conversión DOCX->PDF no produjo archivo PDF. stdout=%s stderr=%s",
            resultado.stdout,
            resultado.stderr,
        )
        return False
    except Exception as exc:
        current_app.logger.warning("No se pudo convertir DOCX a PDF: %s", exc)
        return False


def _obtener_datos_plaza_para_informe(id_plaza: int) -> dict:
    """
    Recupera los datos necesarios para construir el informe de adjudicación.
    """
    datos = ejecutar_query(
        """
        SELECT
            p.idtbl_plazas,
            p.codigo_plazas,
            p.fila,
            u.idtbl_usuarios,
            u.numero_cuenta,
            pr.Nombre_Razon_Social AS nombre_proveedor,
            pr.NIF AS nif_proveedor
        FROM tbl_plazas AS p
        LEFT JOIN tbl_usuarios AS u
            ON p.idtbl_usuarios = u.idtbl_usuarios
        LEFT JOIN bd_tbl_comunes.tbl_proveedores AS pr
            ON u.idtbl_proveedores = pr.Idtbl_proveedores
        WHERE p.idtbl_plazas = %s
        """,
        params=(id_plaza,),
        nombre_bd="parquin_camiones",
    )

    if not datos:
        raise ValueError(f"No se encontraron datos para la plaza {id_plaza}")

    return datos[0]


def _crear_documento_adjudicacion(datos: dict, fecha_efectos: datetime.date) -> Path:
    """
    Construye el documento Word de adjudicación y lo guarda en PADRON_DIR.
    """
    codigo_plaza = datos.get("codigo_plazas")
    nombre_proveedor = datos.get("nombre_proveedor") or "________________"
    nif_proveedor = datos.get("nif_proveedor") or "__________"
    hoy = datetime.date.today()

    nombre_fichero = f"Adjudicacion_Plaza_{codigo_plaza}_Rio_Torio.docx"
    ruta_docx = PADRON_DIR / nombre_fichero

    document = Document()
    section = document.sections[0]
    section.left_margin = Cm(3)

    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_paragraph.add_run("EXCMO. AYUNTAMIENTO DE ÁVILA\nPolicía Local\n")

    p_exp = document.add_paragraph()
    p_exp.add_run("EXPT. xxxx/yyyy").bold = True

    p_asunto = document.add_paragraph()
    p_asunto.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_asunto.add_run(
        f"ASUNTO: Adjudicación de la plaza nº {codigo_plaza} de parking de la calle Río Torío."
    ).bold = True

    document.add_paragraph()

    p_cuerpo = document.add_paragraph()
    p_cuerpo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_cuerpo.paragraph_format.first_line_indent = Cm(2)
    p_cuerpo.add_run("Vistas las solicitudes formuladas por ")
    p_cuerpo.add_run(nombre_proveedor).bold = True
    p_cuerpo.add_run(", con D.N.I. ")
    p_cuerpo.add_run(nif_proveedor).bold = True
    p_cuerpo.add_run(
        ", se comunica que le ha sido adjudicada la plaza nº "
        f"{codigo_plaza}, con fecha de efectos {fecha_efectos.strftime('%d/%m/%Y')}."
    )

    document.add_paragraph()

    p_cierre = document.add_paragraph()
    p_cierre.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_cierre.paragraph_format.first_line_indent = Cm(2)
    p_cierre.add_run("Lo que comunico a Vd. para su conocimiento y efectos oportunos.")

    document.add_paragraph()

    p_fecha = document.add_paragraph()
    p_fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_fecha.add_run(f"Ávila, a {hoy.strftime('%d/%m/%Y')}.")

    document.add_paragraph()

    p_cargo = document.add_paragraph()
    p_cargo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cargo.add_run("El Jefe de Policía Local")

    p_nombre = document.add_paragraph()
    p_nombre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_nombre.add_run("D. Carlos Blanco Rubio")

    p_dest = document.add_paragraph()
    p_dest.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_dest.add_run("EXCMO. AYUNTAMIENTO DE ÁVILA\nGESTIÓN TRIBUTARIA")

    document.save(ruta_docx)
    return ruta_docx


def _generar_pdf_y_docx_adjudicacion(id_plaza: int) -> tuple[Path, Path]:
    """
    Genera el DOCX y su copia PDF para una plaza concreta.
    """
    datos = _obtener_datos_plaza_para_informe(id_plaza)
    hoy = datetime.date.today()
    fecha_efectos = _mes_siguiente(hoy)

    ruta_docx = _crear_documento_adjudicacion(datos, fecha_efectos)
    ruta_pdf = ruta_docx.with_suffix(".pdf")
    _generar_pdf_desde_docx(ruta_docx, ruta_pdf)

    return ruta_docx, ruta_pdf


# =============================================================================
# VISTA PRINCIPAL
# =============================================================================


@btn_rio_torio_asignar_plaza_bp.route("/asignar_plaza", methods=["GET", "POST"])
@login_required
@rol_required("gestor", "super_admin")
def btn_rio_torio_asignar_plaza():
    """
    Vista principal para asignar o liberar plazas.

    Flujo:
    - GET: carga formularios y tablas.
    - POST: aplica cambios, registra histórico y genera informe.
    """
    if request.method == "POST":
        # ---------------------------------------------------------------------
        # RECUPERACIÓN DE PARÁMETROS DEL FORMULARIO
        # ---------------------------------------------------------------------
        # id_usuario: usuario al que se asignan plazas (nuevo ocupante).
        id_usuario = request.form.get("idtbl_usuarios")

        # Listas de plazas:
        # - plazas_asignar: plazas seleccionadas para asignar al usuario.
        # - plazas_quitar: plazas que se van a liberar del usuario actual.
        plazas_asignar = request.form.getlist("plazas_seleccionadas")
        plazas_quitar = request.form.getlist("plazas_quitar_usuario")

        # Fechas de asignación y liberación.
        fecha_asignacion_str = request.form.get("fecha_asignacion")
        fecha_liberacion_str = request.form.get("fecha_liberacion")

        # Número de expediente para la asignación (puede venir de la solicitud).
        numero_expediente_asignacion = request.form.get("numero_expediente_asignacion")

        # Solicitud de plaza que se va a utilizar para el alta (seleccionada en el formulario).
        id_solicitud_plaza = request.form.get("idtbl_solicitudes_plazas")

        # Flag para generar informe (DOCX + PDF).
        generar_informe = request.form.get("generar_informe") == "1"

        # ---------------------------------------------------------------------
        # PARSEO DE FECHAS
        # ---------------------------------------------------------------------
        # Si no se indica fecha_asignacion/fecha_liberacion, se usan valores actuales.
        fecha_asignacion = (
            datetime.datetime.strptime(fecha_asignacion_str, "%Y-%m-%d")
            if fecha_asignacion_str
            else datetime.datetime.now()
        )
        fecha_liberacion = (
            datetime.datetime.strptime(fecha_liberacion_str, "%Y-%m-%d")
            if fecha_liberacion_str
            else datetime.datetime.now()
        )

        # ---------------------------------------------------------------------
        # NORMALIZACIÓN DE CONJUNTOS DE PLAZAS
        # ---------------------------------------------------------------------
        # set_quitar: plazas que se van a liberar.
        # set_asignar: plazas que se van a asignar al usuario (excluyendo las que se quitan).
        set_quitar = set(plazas_quitar)
        set_asignar = set(plazas_asignar) - set_quitar

        try:
            # -----------------------------------------------------------------
            # APERTURA DE CONEXIÓN Y INICIO DE TRANSACCIÓN
            # -----------------------------------------------------------------
            conn = get_connection("parquin_camiones")
            cursor = conn.cursor(dictionary=True)

            # Se inicia una transacción explícita para asegurar atomicidad:
            # o se aplican todos los cambios (bajas + altas + solicitudes + histórico),
            # o no se aplica ninguno.
            conn.start_transaction()

            # =================================================================
            # 1) BAJA DE PLAZAS: CERRAR HISTÓRICO ACTIVO Y LIBERAR PLAZAS
            # =================================================================
            if set_quitar:
                ids_quitar = [int(pid) for pid in set_quitar]
                placeholders_quitar = ", ".join(["%s"] * len(ids_quitar))

                # --------------------------------------------------------------
                # 1.1) Recuperar datos de las plazas que se van a quitar:
                #      - plaza actual (idtbl_plazas)
                #      - usuario y proveedor actual
                #      - fecha_inicio y expediente para poder cerrar histórico.
                # --------------------------------------------------------------
                cursor.execute(
                    f"""
                    SELECT
                        p.idtbl_plazas,
                        p.fecha_inicio,
                        p.numero_expediente,
                        p.exp_solicitud_fin,
                        p.idtbl_usuarios,
                        u.idtbl_proveedores
                    FROM tbl_plazas AS p
                    LEFT JOIN tbl_usuarios AS u
                        ON p.idtbl_usuarios = u.idtbl_usuarios
                    WHERE p.idtbl_plazas IN ({placeholders_quitar})
                      AND p.idtbl_usuarios IS NOT NULL
                    FOR UPDATE
                    """,
                    tuple(ids_quitar),
                )
                plazas_a_quitar = cursor.fetchall()

                # --------------------------------------------------------------
                # 1.2) Cerrar histórico activo por plaza+proveedor:
                #      Se busca el registro de histórico de esa plaza y proveedor
                #      con fecha_fin IS NULL y se actualiza con la fecha de baja.
                #      No se crea un registro nuevo; se actualiza el existente.
                # --------------------------------------------------------------
                for fila in plazas_a_quitar:
                    id_plaza = fila["idtbl_plazas"]
                    id_usuario_actual = fila["idtbl_usuarios"]
                    id_proveedor_actual = fila["idtbl_proveedores"]
                    exp_solicitud_fin = fila["exp_solicitud_fin"]

                    if id_proveedor_actual is None:
                        # Si la plaza no tiene proveedor asociado, no se puede
                        # cerrar histórico por proveedor; se deja constancia en log.
                        current_app.logger.warning(
                            "Plaza %s sin proveedor asociado al hacer baja; no se cierra histórico",
                            id_plaza,
                        )
                        continue

                    # Cerrar el histórico activo de esta plaza y proveedor.
                    cursor.execute(
                        """
                        UPDATE tbl_historico_plazas AS h
                        SET
                            h.fecha_fin = %s,
                            h.exp_solicitud_fin = %s,
                            h.idtbl_usuarios = %s
                        WHERE h.idtbl_plazas = %s
                          AND h.idtbl_proveedores = %s
                          AND h.fecha_fin IS NULL
                        """,
                        (
                            fecha_liberacion,
                            exp_solicitud_fin,
                            id_usuario_actual,
                            id_plaza,
                            id_proveedor_actual,
                        ),
                    )

                # --------------------------------------------------------------
                # 1.3) Liberar las plazas en tbl_plazas:
                #      Se borra la relación con el usuario y se limpian campos
                #      de fecha_inicio y expediente para dejarla disponible.
                # --------------------------------------------------------------
                cursor.execute(
                    f"""
                    UPDATE tbl_plazas
                    SET
                        idtbl_usuarios = NULL,
                        fecha_inicio = NULL,
                        fecha_fin = NULL,
                        numero_expediente = NULL,
                        exp_solicitud = NULL,
                        exp_solicitud_fin = NULL
                    WHERE idtbl_plazas IN ({placeholders_quitar})
                    """,
                    tuple(ids_quitar),
                )

            # =================================================================
            # 2) ALTA DE PLAZAS: NUEVO HISTÓRICO + BLOQUEO + SOLICITUD APROBADA
            # =================================================================
            if set_asignar and id_usuario:
                ids_asignar = [int(pid) for pid in set_asignar]
                placeholders_asignar = ", ".join(["%s"] * len(ids_asignar))

                # --------------------------------------------------------------
                # 2.1) Recuperar usuario asignado y su proveedor/forma_pago:
                # --------------------------------------------------------------
                cursor.execute(
                    """
                    SELECT
                        idtbl_usuarios,
                        idtbl_proveedores,
                        forma_pago
                    FROM tbl_usuarios
                    WHERE idtbl_usuarios = %s
                    """,
                    (id_usuario,),
                )
                usuario_asignado = cursor.fetchone()

                if not usuario_asignado:
                    raise ValueError(
                        "No se ha encontrado el usuario seleccionado para la asignación."
                    )

                id_proveedor_nuevo = usuario_asignado["idtbl_proveedores"]
                forma_pago_usuario = usuario_asignado["forma_pago"]

                # --------------------------------------------------------------
                # 2.2) Recuperar la solicitud seleccionada (si se ha indicado):
                #      Esta solicitud enlazará la alta con su expediente y forma de pago.
                # --------------------------------------------------------------
                solicitud = None
                n_expediente_solicitud = None
                forma_pago_solicitud = None

                if id_solicitud_plaza:
                    cursor.execute(
                        """
                        SELECT
                            idtbl_solicitudes_plazas,
                            idtbl_usuarios,
                            idtbl_plazas,
                            n_expediente,
                            fecha_solicitud,
                            forma_pago,
                            estado
                        FROM tbl_solicitudes_plazas
                        WHERE idtbl_solicitudes_plazas = %s
                        FOR UPDATE
                        """,
                        (id_solicitud_plaza,),
                    )
                    solicitud = cursor.fetchone()

                    if not solicitud:
                        raise ValueError("La solicitud seleccionada no existe.")
                    if solicitud["estado"] != "pendiente":
                        raise ValueError(
                            "La solicitud seleccionada no está en estado pendiente."
                        )

                    # Se usan expediente y forma de pago de la solicitud para el histórico.
                    n_expediente_solicitud = solicitud["n_expediente"]
                    forma_pago_solicitud = solicitud["forma_pago"]

                # --------------------------------------------------------------
                # 2.3) Recuperar plazas actuales antes de asignar:
                #      Si la plaza ya estaba ocupada por otro proveedor, se cierra
                #      el histórico anterior (cambio de proveedor).
                # --------------------------------------------------------------
                cursor.execute(
                    f"""
                    SELECT
                        p.idtbl_plazas,
                        p.fecha_inicio,
                        p.numero_expediente,
                        p.exp_solicitud,
                        p.idtbl_usuarios AS idtbl_usuarios_anterior,
                        u.idtbl_proveedores AS idtbl_proveedores_anterior
                    FROM tbl_plazas AS p
                    LEFT JOIN tbl_usuarios AS u
                        ON p.idtbl_usuarios = u.idtbl_usuarios
                    WHERE p.idtbl_plazas IN ({placeholders_asignar})
                    FOR UPDATE
                    """,
                    tuple(ids_asignar),
                )
                plazas_previas = cursor.fetchall()

                for fila_plaza in plazas_previas:
                    id_plaza = fila_plaza["idtbl_plazas"]
                    id_usuario_anterior = fila_plaza["idtbl_usuarios_anterior"]
                    id_proveedor_anterior = fila_plaza["idtbl_proveedores_anterior"]
                    fecha_inicio_anterior = fila_plaza["fecha_inicio"]
                    exp_solicitud_anterior = fila_plaza["exp_solicitud"]

                    # ----------------------------------------------------------
                    # 2.3.1) Si la plaza estaba ocupada por otro proveedor, se
                    #        cierra su histórico activo (cambio de proveedor).
                    # ----------------------------------------------------------
                    if id_proveedor_anterior is not None:
                        cursor.execute(
                            """
                            UPDATE tbl_historico_plazas AS h
                            SET
                                h.fecha_fin = %s,
                                h.exp_solicitud_fin = %s,
                                h.idtbl_usuarios = %s
                            WHERE h.idtbl_plazas = %s
                              AND h.idtbl_proveedores = %s
                              AND h.fecha_fin IS NULL
                            """,
                            (
                                fecha_asignacion,
                                exp_solicitud_anterior,
                                id_usuario_anterior,
                                id_plaza,
                                id_proveedor_anterior,
                            ),
                        )

                    # ----------------------------------------------------------
                    # 2.3.2) Actualizar tbl_plazas con el nuevo usuario,
                    #        fecha de inicio y expediente (de solicitud o del form).
                    # ----------------------------------------------------------
                    if set_asignar and id_usuario and not id_solicitud_plaza:
                        raise ValueError(
                            "Para asignar plazas es obligatorio seleccionar una solicitud."
                        )

                    numero_expediente_final = (
                        n_expediente_solicitud or numero_expediente_asignacion
                    )
                    exp_solicitud_final = (
                        n_expediente_solicitud or fila_plaza["exp_solicitud"]
                    )

                    cursor.execute(
                        """
                        UPDATE tbl_plazas
                        SET
                            idtbl_usuarios = %s,
                            fecha_inicio = %s,
                            fecha_fin = NULL,
                            numero_expediente = %s,
                            exp_solicitud = %s,
                            exp_solicitud_fin = NULL
                        WHERE idtbl_plazas = %s
                        """,
                        (
                            id_usuario,
                            fecha_asignacion,
                            numero_expediente_final,
                            exp_solicitud_final,
                            id_plaza,
                        ),
                    )

                    # ----------------------------------------------------------
                    # 2.3.3) Insertar nuevo histórico de alta para la plaza:
                    #        - fecha_inicio = fecha_asignacion
                    #        - fecha_fin = NULL (registro activo)
                    #        - exp_solicitud_inicio = n_expediente_solicitud
                    #        - forma_pago = de solicitud o del usuario
                    # ----------------------------------------------------------
                    forma_pago_final = forma_pago_solicitud or forma_pago_usuario

                    cursor.execute(
                        """
                        INSERT INTO tbl_historico_plazas (
                            idtbl_plazas,
                            idtbl_proveedores,
                            fecha_inicio,
                            fecha_fin,
                            exp_solicitud_fin,
                            forma_pago,
                            observaciones,
                            idtbl_usuarios,
                            exp_solicitud_inicio,
                            exp_solicitud_cambio
                        )
                        VALUES (
                            %s,  -- idtbl_plazas
                            %s,  -- idtbl_proveedores nuevo
                            %s,  -- fecha_inicio (alta)
                            NULL, -- fecha_fin (activo)
                            NULL, -- exp_solicitud_fin (se rellenará al hacer baja)
                            %s,  -- forma_pago
                            NULL, -- observaciones
                            %s,  -- idtbl_usuarios nuevo
                            %s,  -- exp_solicitud_inicio (n_expediente solicitud)
                            %s   -- exp_solicitud_cambio (si venía de otro proveedor)
                        )
                        """,
                        (
                            id_plaza,
                            id_proveedor_nuevo,
                            fecha_asignacion,
                            forma_pago_final,
                            int(id_usuario),
                            n_expediente_solicitud,
                            exp_solicitud_anterior,
                        ),
                    )

                # --------------------------------------------------------------
                # 2.4) Marcar la solicitud de plaza como aprobada (si existe):
                #      Se registra la fecha_aprobacion y se actualiza estado.
                # --------------------------------------------------------------
                if solicitud:
                    cursor.execute(
                        """
                        UPDATE tbl_solicitudes_plazas
                        SET
                            estado = 'aprobada',
                            fecha_aprobacion = %s,
                            updated_at = NOW()
                        WHERE idtbl_solicitudes_plazas = %s
                        """,
                        (
                            fecha_asignacion,
                            solicitud["idtbl_solicitudes_plazas"],
                        ),
                    )

            # =================================================================
            # 3) COMMIT DE LA TRANSACCIÓN
            # =================================================================
            conn.commit()
            flash("Plazas actualizadas correctamente", "success")

        except Exception as exc:
            # En caso de error, se deshacen todos los cambios de la transacción.
            conn.rollback()
            current_app.logger.exception("Error actualizando plazas")
            flash(f"Error al actualizar plazas: {exc}", "danger")

        finally:
            # Cierre de cursor y conexión.
            cursor.close()
            conn.close()

        # =====================================================================
        # 4) GENERACIÓN DE INFORME (DOCX+PDF) SI SE HA SELECCIONADO
        # =====================================================================
        if generar_informe:
            try:
                # Se intenta generar informe para la primera plaza afectada
                # (priorizando altas; si no, bajas).
                id_plaza_informe = int(next(iter(set_asignar or set_quitar)))
                ruta_docx, ruta_pdf = _generar_pdf_y_docx_adjudicacion(id_plaza_informe)

                flash(f"Informe Word generado: {ruta_docx.name}", "success")
                flash(f"Copia PDF guardada: {ruta_pdf.name}", "success")

                return send_file(
                    ruta_docx,
                    as_attachment=True,
                    download_name=ruta_docx.name,
                    mimetype=(
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    ),
                )
            except Exception as exc:
                current_app.logger.exception("Error generando informe de adjudicación")
                flash(f"Error generando el informe: {exc}", "danger")

        return redirect(
            url_for("btn_rio_torio_asignar_plaza_bp.btn_rio_torio_asignar_plaza")
        )

    usuarios_parquin, plazas, sort, dir_param = _cargar_datos_parquin_rio_torio()
    return render_template(
        "parquin/rio_torio/rio_torio_asignar_plaza.html",
        usuarios_parquin=usuarios_parquin,
        plazas=plazas,
        sort=sort,
        dir_param=dir_param,
    )


# =============================================================================
# VISTA INFORME INDIVIDUAL
# =============================================================================


@btn_rio_torio_asignar_plaza_bp.route(
    "/informe_adjudicacion/<int:id_plaza>", methods=["GET"]
)
@login_required
@rol_required("gestor", "super_admin")
def rio_torio_informe_adjudicacion(id_plaza: int):
    """
    Genera informe Word + PDF y devuelve el documento Word.
    """
    try:
        ruta_docx, ruta_pdf = _generar_pdf_y_docx_adjudicacion(id_plaza)
        flash(f"Informe Word generado: {ruta_docx.name}", "success")
        flash(f"Copia PDF guardada: {ruta_pdf.name}", "success")
        return send_file(
            ruta_docx,
            as_attachment=True,
            download_name=ruta_docx.name,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    except Exception as exc:
        current_app.logger.exception("Error generando informe de adjudicación")
        flash(f"Error generando el informe de adjudicación: {exc}", "danger")
        return redirect(
            url_for("btn_rio_torio_asignar_plaza_bp.btn_rio_torio_asignar_plaza")
        )

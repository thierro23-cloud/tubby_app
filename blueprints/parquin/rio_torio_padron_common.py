# =============================================================================
# 🧾 PADRÓN RÍO TORÍO · LÓGICA COMÚN
# =============================================================================
from __future__ import annotations

from collections import defaultdict
from datetime import date, timedelta
from pathlib import Path
import calendar

from flask import current_app
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

from db import ejecutar_query


BASE_DIR = Path(__file__).resolve().parent.parent


def localizar_padron(base_dir: Path) -> Path | None:
    for p in base_dir.rglob("padron"):
        if p.is_dir():
            return p
    return None


PADRON_DIR = localizar_padron(BASE_DIR)
if PADRON_DIR is None:
    PADRON_DIR = BASE_DIR / "parquin" / "padron"

RUTA_DESTINO = PADRON_DIR

MESES_ES = [
    "",
    "enero",
    "febrero",
    "marzo",
    "abril",
    "mayo",
    "junio",
    "julio",
    "agosto",
    "septiembre",
    "octubre",
    "noviembre",
    "diciembre",
]


def obtener_periodo_padron(fecha_base: date | None = None) -> dict:
    """
    Periodo automático del padrón: mes anterior al primer laborable del mes siguiente.
    """
    if fecha_base is None:
        fecha_base = date.today()

    if fecha_base.month == 12:
        anio_siguiente = fecha_base.year + 1
        mes_siguiente = 1
    else:
        anio_siguiente = fecha_base.year
        mes_siguiente = fecha_base.month + 1

    primer_dia_mes_siguiente = date(anio_siguiente, mes_siguiente, 1)
    primer_laborable = primer_dia_mes_siguiente
    while primer_laborable.weekday() >= 5:
        primer_laborable += timedelta(days=1)

    if mes_siguiente == 1:
        anio_padron = anio_siguiente - 1
        mes_padron = 12
    else:
        anio_padron = anio_siguiente
        mes_padron = mes_siguiente - 1

    inicio_mes = date(anio_padron, mes_padron, 1)
    _, ultimo_dia = calendar.monthrange(anio_padron, mes_padron)
    fin_mes = date(anio_padron, mes_padron, ultimo_dia)

    return {
        "inicio_mes": inicio_mes,
        "fin_mes": fin_mes,
        "anio": anio_padron,
        "mes_num": mes_padron,
        "mes_nombre": MESES_ES[mes_padron],
        "primer_laborable_mes_siguiente": primer_laborable,
    }


def periodo_mes_anio(mes: int, anio: int) -> tuple[date, date]:
    inicio = date(anio, mes, 1)
    _, ultimo_dia = calendar.monthrange(anio, mes)
    fin = date(anio, mes, ultimo_dia)
    return inicio, fin


def etiqueta_periodo(inicio: date, fin: date) -> str:
    if inicio.day == 1 and inicio.year == fin.year and inicio.month == fin.month:
        _, ultimo_dia = calendar.monthrange(inicio.year, inicio.month)
        if fin.day == ultimo_dia:
            return f"{MESES_ES[inicio.month].capitalize()} de {inicio.year}"
    return f"del {inicio.strftime('%d/%m/%Y')} al {fin.strftime('%d/%m/%Y')}"


def nombre_fichero_periodo(inicio: date, fin: date, sufijo: str = "") -> str:
    if inicio.day == 1 and inicio.year == fin.year and inicio.month == fin.month:
        _, ultimo_dia = calendar.monthrange(inicio.year, inicio.month)
        if fin.day == ultimo_dia:
            base = f"{MESES_ES[inicio.month].capitalize()}_{inicio.year}_Padron_de_parking_Rio_Torio"
        else:
            base = f"{inicio.strftime('%Y%m%d')}_{fin.strftime('%Y%m%d')}_Padron_de_parking_Rio_Torio"
    else:
        base = f"{inicio.strftime('%Y%m%d')}_{fin.strftime('%Y%m%d')}_Padron_de_parking_Rio_Torio"

    if sufijo:
        base = f"{base}_{sufijo}"
    return f"{base}.docx"


def obtener_historico_periodo(inicio_mes: date, fin_mes: date) -> list[dict]:
    sql = (
        "SELECT "
        "    h.idtbl_historico_plazas, "
        "    h.idtbl_plazas, "
        "    h.idtbl_proveedores, "
        "    h.fecha_inicio, "
        "    h.fecha_fin, "
        "    h.exp_solicitud_fin, "
        "    h.forma_pago, "
        "    pl.codigo_plazas, "
        "    pl.fila, "
        "    CONCAT_WS(' ', p.apellidos, p.nombre) AS nombre_proveedor, "
        "    p.NIF AS nif_proveedor "
        "FROM parquin_camiones.tbl_historico_plazas AS h "
        "JOIN parquin_camiones.tbl_plazas AS pl "
        "  ON h.idtbl_plazas = pl.idtbl_plazas "
        "JOIN bd_tbl_comunes.tbl_proveedores AS p "
        "  ON h.idtbl_proveedores = p.Idtbl_proveedores "
        "WHERE "
        "    h.fecha_inicio <= %s "
        "    AND (h.fecha_fin IS NULL OR h.fecha_fin >= %s) "
        "ORDER BY p.apellidos, p.nombre, pl.codigo_plazas, h.fecha_inicio"
    )
    return ejecutar_query(sql, params=(fin_mes, inicio_mes), nombre_bd="parquin_camiones")


def construir_padron_principal(filas: list[dict]) -> list[dict]:
    por_proveedor: dict[int, dict] = {}

    for f in filas:
        id_prov = f["idtbl_proveedores"]
        if id_prov not in por_proveedor:
            por_proveedor[id_prov] = {
                "nombre_proveedor": f["nombre_proveedor"],
                "nif": f["nif_proveedor"],
                "plazas": set(),
                "forma_pago": (f.get("forma_pago") or ""),
            }
        por_proveedor[id_prov]["plazas"].add(f["codigo_plazas"])

    resultado = []
    for prov in por_proveedor.values():
        plazas_ordenadas = sorted(prov["plazas"], key=str)
        resultado.append(
            {
                "nombre_proveedor": prov["nombre_proveedor"],
                "nif": prov["nif"],
                "plazas": plazas_ordenadas,
                "total_plazas": len(plazas_ordenadas),
                "forma_pago": prov["forma_pago"],
            }
        )
    return resultado


def construir_variaciones(filas: list[dict], inicio_mes: date, fin_mes: date) -> list[dict]:
    variaciones: list[dict] = []
    por_prov: dict[int, list[dict]] = defaultdict(list)

    for f in filas:
        por_prov[f["idtbl_proveedores"]].append(f)

    for id_prov, registros in por_prov.items():
        registros.sort(key=lambda r: (r["fecha_inicio"], r["idtbl_plazas"]))

        bajas = []
        altas = []
        for r in registros:
            fi = r["fecha_inicio"]
            ff = r["fecha_fin"]
            if inicio_mes <= fi <= fin_mes:
                altas.append(r)
            if ff is not None and inicio_mes <= ff <= fin_mes:
                bajas.append(r)

        usadas_bajas: set[int] = set()
        usadas_altas: set[int] = set()

        for baja in bajas:
            ff = baja["fecha_fin"]
            if ff is None:
                continue
            dia_siguiente = ff + timedelta(days=1)

            for idx_alta, alta in enumerate(altas):
                if idx_alta in usadas_altas:
                    continue
                if (
                    alta["fecha_inicio"] == dia_siguiente
                    and alta["idtbl_plazas"] != baja["idtbl_plazas"]
                    and alta["idtbl_proveedores"] == id_prov
                ):
                    variaciones.append(
                        {
                            "nombre_proveedor": baja["nombre_proveedor"],
                            "nif": baja["nif_proveedor"],
                            "numero_plaza": f"{baja['codigo_plazas']} -> {alta['codigo_plazas']}",
                            "tipo_cambio": "Cambio de plaza",
                            "fecha": ff,
                            "forma_pago": baja.get("forma_pago"),
                        }
                    )
                    usadas_bajas.add(baja["idtbl_historico_plazas"])
                    usadas_altas.add(idx_alta)
                    break

        for baja in bajas:
            if baja["idtbl_historico_plazas"] in usadas_bajas:
                continue
            variaciones.append(
                {
                    "nombre_proveedor": baja["nombre_proveedor"],
                    "nif": baja["nif_proveedor"],
                    "numero_plaza": baja["codigo_plazas"],
                    "tipo_cambio": "BAJA",
                    "fecha": baja["fecha_fin"],
                    "forma_pago": baja.get("forma_pago"),
                }
            )

        por_plaza: dict[int, list[dict]] = defaultdict(list)
        for r in registros:
            por_plaza[r["idtbl_plazas"]].append(r)

        for regs_plaza in por_plaza.values():
            regs_plaza.sort(key=lambda r: r["fecha_inicio"])
            forma_anterior = None
            for r in regs_plaza:
                forma_actual = r.get("forma_pago") or ""
                fi = r["fecha_inicio"]
                if forma_anterior is not None and forma_actual != forma_anterior:
                    if inicio_mes <= fi <= fin_mes:
                        variaciones.append(
                            {
                                "nombre_proveedor": r["nombre_proveedor"],
                                "nif": r["nif_proveedor"],
                                "numero_plaza": r["codigo_plazas"],
                                "tipo_cambio": "Forma de pago",
                                "fecha": fi,
                                "forma_pago": forma_actual,
                            }
                        )
                forma_anterior = forma_actual

    variaciones.sort(key=lambda v: (v["nombre_proveedor"], v["fecha"], v["numero_plaza"]))
    return variaciones


def obtener_resumen_plazas_rio_torio(inicio_periodo: date, fin_periodo: date) -> dict:
    """
    Resumen calculado SIEMPRE desde tbl_historico_plazas.

    - totales: plazas distintas que aparecen en el histórico.
    - ocupadas: plazas distintas activas dentro del periodo solicitado.
    - libres: diferencia entre totales históricos y ocupadas del periodo.

    No consulta tbl_plazas.idtbl_usuarios ni el estado actual de las plazas,
    para que el padrón sea coherente con el periodo histórico elegido.
    """
    sql = (
        "SELECT "
        "    COUNT(DISTINCT h.idtbl_plazas) AS total, "
        "    COUNT(DISTINCT CASE "
        "        WHEN h.fecha_inicio <= %s "
        "         AND (h.fecha_fin IS NULL OR h.fecha_fin >= %s) "
        "        THEN h.idtbl_plazas "
        "    END) AS ocupadas "
        "FROM parquin_camiones.tbl_historico_plazas AS h"
    )

    filas = ejecutar_query(
        sql,
        params=(fin_periodo, inicio_periodo),
        nombre_bd="parquin_camiones",
    )

    row = filas[0] if filas else {"total": 0, "ocupadas": 0}
    total = row.get("total") or 0
    ocupadas = row.get("ocupadas") or 0
    libres = max(total - ocupadas, 0)

    return {
        "totales": total,
        "libres": libres,
        "ocupadas": ocupadas,
    }


def contar_formas_pago(filas: list[dict]) -> dict:
    plazas = {
        "mensual": set(),
        "trimestral": set(),
        "semestral": set(),
        "anual": set(),
        "otros": set(),
    }

    for f in filas:
        forma = (f.get("forma_pago") or "").strip().lower()
        plaza = f.get("codigo_plazas")
        if not plaza:
            continue
        if forma in ("mensual", "m"):
            plazas["mensual"].add(plaza)
        elif forma in ("trimestral", "t"):
            plazas["trimestral"].add(plaza)
        elif forma in ("semestral", "s"):
            plazas["semestral"].add(plaza)
        elif forma in ("anual", "a"):
            plazas["anual"].add(plaza)
        else:
            plazas["otros"].add(plaza)

    return {k: len(v) for k, v in plazas.items()}


def generar_informe_word_periodo(inicio_periodo: date, fin_periodo: date, sufijo: str = "") -> Path:
    """
    Genera exactamente el mismo informe para cualquier origen:
      - Manual mes/año.
      - Manual rango.
      - Automático.
    Lo único que cambia es el periodo recibido.
    """
    RUTA_DESTINO.mkdir(parents=True, exist_ok=True)
    ruta_completa = RUTA_DESTINO / nombre_fichero_periodo(inicio_periodo, fin_periodo, sufijo=sufijo)
    periodo_texto = etiqueta_periodo(inicio_periodo, fin_periodo)

    filas = obtener_historico_periodo(inicio_periodo, fin_periodo)
    padron_principal = construir_padron_principal(filas)
    variaciones = construir_variaciones(filas, inicio_periodo, fin_periodo)
    contadores = contar_formas_pago(filas)
    resumen = obtener_resumen_plazas_rio_torio(inicio_periodo, fin_periodo)

    document = Document()
    section = document.sections[0]
    hoy = date.today()

    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_paragraph.add_run("Jefatura de Policía Local\nAdministración\n")
    header_paragraph.add_run(
        f"Plazas totales: {resumen['totales']} | "
        f"Libres: {resumen['libres']} | "
        f"Ocupadas: {resumen['ocupadas']}"
    )

    section.left_margin = Cm(3)

    p_asunto = document.add_paragraph()
    p_asunto.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_asunto.add_run(
        "ASUNTO: Relación de abonados y plazas del parking de camiones de la "
        "Calle Río Torío (formato conjunto)."
    ).bold = True

    document.add_paragraph()

    p_interesado = document.add_paragraph()
    p_interesado.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_interesado.add_run("Interesado: Gestión Tributaria.")

    document.add_paragraph()

    p_comunico = document.add_paragraph()
    p_comunico.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_comunico.paragraph_format.first_line_indent = Cm(2)
    p_comunico.add_run(
        f"Comunico a Vd. que durante el periodo {periodo_texto} "
        "la relación de abonados y sus plazas es la siguiente:"
    )

    document.add_paragraph()

    p_contadores = document.add_paragraph()
    p_contadores.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_contadores.paragraph_format.first_line_indent = Cm(2)
    texto_contadores = (
        f"total pago mensual = {contadores['mensual']} | "
        f"trimestral = {contadores['trimestral']} | "
        f"semestral = {contadores['semestral']} | "
        f"anual = {contadores['anual']}"
    )
    if contadores["otros"]:
        texto_contadores += f" | otros = {contadores['otros']}"
    p_contadores.add_run(texto_contadores).bold = True

    document.add_paragraph()

    if padron_principal:
        table = document.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Apellidos y nombre"
        hdr_cells[1].text = "NIF"
        hdr_cells[2].text = "Forma de pago"
        hdr_cells[3].text = "Plazas"
        hdr_cells[4].text = "Total plazas"

        for a in padron_principal:
            row = table.add_row().cells
            row[0].text = str(a.get("nombre_proveedor", ""))
            row[1].text = str(a.get("nif", ""))
            row[2].text = (a.get("forma_pago") or "").upper()
            row[3].text = ", ".join(a.get("plazas", []))
            row[4].text = str(a.get("total_plazas", 0) or 0)
    else:
        p_sin = document.add_paragraph()
        p_sin.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_sin.add_run("Actualmente no existen abonados registrados.").italic = True

    document.add_paragraph()

    if variaciones:
        p_var_titulo = document.add_paragraph()
        p_var_titulo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_var_titulo.paragraph_format.first_line_indent = Cm(2)
        p_var_titulo.add_run(
            "Variaciones (bajas, cambios de plaza, cambios de forma de pago):"
        ).bold = True

        table_var = document.add_table(rows=1, cols=5)
        table_var.style = "Table Grid"
        hdrv = table_var.rows[0].cells
        hdrv[0].text = "Fecha"
        hdrv[1].text = "Proveedor"
        hdrv[2].text = "NIF"
        hdrv[3].text = "Forma de pago"
        hdrv[4].text = "Detalle"

        for v in variaciones:
            rowv = table_var.add_row().cells
            fv = v.get("fecha")
            rowv[0].text = fv.strftime("%d/%m/%Y") if fv else ""
            rowv[1].text = v.get("nombre_proveedor", "")
            rowv[2].text = v.get("nif", "")
            rowv[3].text = (v.get("forma_pago") or "").upper()
            rowv[4].text = f"{v.get('tipo_cambio', '')}: {v.get('numero_plaza', '')}"
    else:
        p_sin_var = document.add_paragraph()
        p_sin_var.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_sin_var.paragraph_format.first_line_indent = Cm(2)
        p_sin_var.add_run("SIN VARIACIONES EN ESTE PERIODO.").bold = True

    document.add_paragraph()

    p_final = document.add_paragraph()
    p_final.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_final.paragraph_format.first_line_indent = Cm(2)
    p_final.add_run("Lo que comunico a Vd., para su conocimiento y efectos oportunos.")

    document.add_paragraph()
    p_firma_fecha = document.add_paragraph()
    p_firma_fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_firma_fecha.add_run(f"En Ávila, a {hoy.strftime('%d/%m/%Y')}.")

    document.add_paragraph()
    p_cargo = document.add_paragraph()
    p_cargo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cargo.add_run("El Jefe de la Policía Local.")

    document.add_paragraph()
    p_nombre = document.add_paragraph()
    p_nombre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_nombre.add_run("D. Carlos Blanco Rubio")

    p_digital = document.add_paragraph()
    p_digital.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_digital.add_run("(documento firmado digitalmente)").italic = True

    document.save(ruta_completa)
    current_app.logger.info("Padrón Río Torío generado: %s", ruta_completa)
    return ruta_completa


def datos_padron_periodo(inicio_periodo: date, fin_periodo: date) -> dict:
    filas = obtener_historico_periodo(inicio_periodo, fin_periodo)
    return {
        "periodo": {
            "inicio": inicio_periodo.isoformat(),
            "fin": fin_periodo.isoformat(),
        },
        "padron_principal": construir_padron_principal(filas),
        "variaciones": construir_variaciones(filas, inicio_periodo, fin_periodo),
        "contadores": contar_formas_pago(filas),
        "resumen_plazas": obtener_resumen_plazas_rio_torio(inicio_periodo, fin_periodo),
    }

from __future__ import annotations

from datetime import date, timedelta
from pathlib import Path

from flask import (
    Blueprint,
    current_app,
    flash,
    redirect,
    url_for,
    session,
    request,
    render_template,
)

from services.helpers import rol_required
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from db import ejecutar_query, get_connection
from flask import send_file

# =============================================================================
# 0️⃣ OBJETIVO GENERAL DEL MÓDULO (VERSIÓN "JUNTOS")
# =============================================================================
# Este módulo implementa un botón independiente dentro del área de Río Torío
# que permite:
#   1. Mostrar un formulario donde el usuario selecciona MES y AÑO.
#   2. Al enviar el formulario, generar un documento Word (.docx) con:
#       - Relación de abonados del parking Río Torío.
#       - Para cada abonado, una fila con:
#           · Nombre / Razón social
#           · NIF
#           · Plazas (códigos separados por comas) en una celda
#           · Nº total de plazas en la celda contigua
#   3. Guardar el archivo generado en una carpeta "padron" del proyecto,
#      con nombre <Mes>_<Año>_Padron_de_parking_Rio_Torio_JUNTOS.docx.
# El acceso a la vista está restringido por rol a "gestor" y "super_admin".
# =============================================================================


# =============================================================================
# 1️⃣ RUTAS DE ARCHIVOS Y NOMBRES DE MESES (COMPARTIDO)
# =============================================================================

BASE_DIR = Path(__file__).resolve().parent.parent


def localizar_padron(base_dir: Path) -> Path | None:
    """
    Busca recursivamente, a partir de base_dir, un directorio llamado 'padron'.

    Si lo encuentra, devuelve la primera ruta que coincida.
    Si no existe ningún directorio con ese nombre, devuelve None.
    """
    for p in base_dir.rglob("padron"):
        if p.is_dir():
            return p
    return None


PADRON_DIR = localizar_padron(BASE_DIR)
if PADRON_DIR is None:
    PADRON_DIR = BASE_DIR / "parquin" / "padron"

RUTA_DESTINO = PADRON_DIR

MESES_ES = [
    "",
    "enero",
    "febrero",
    "marzo",
    "abril",
    "mayo",
    "junio",
    "julio",
    "agosto",
    "septiembre",
    "octubre",
    "noviembre",
    "diciembre",
]


# =============================================================================
# 2️⃣ BLUEPRINT DEL BOTÓN DE PADRÓN RÍO TORÍO (VERSIÓN JUNTOS)
# =============================================================================

btn_rio_torio_padron_juntos_bp = Blueprint(
    "btn_rio_torio_padron_juntos_bp",
    __name__,
    url_prefix="/parquin/rio_torio",
)


# =============================================================================
# 3️⃣ HELPERS SQL: ABONADOS Y PLAZAS (VERSIÓN JUNTOS)
# =============================================================================


def _obtener_abonados_y_plazas_juntos(
    fecha_inicio_periodo: date, fecha_fin_periodo: date
):
    """
    Obtiene la información de abonados y sus plazas actuales para la versión
    'juntos', donde se mostrará por abonado:

      - Nombre / Razón social
      - NIF
      - Plazas: lista de códigos de plaza (1, 2, 3, ...)
      - Total de plazas: número de plazas

    Utiliza:
      - parquin_camiones.tbl_usuarios (relación usuario-proveedor).
      - bd_tbl_comunes.tbl_proveedores (nombre y nif).
      - parquin_camiones.tbl_plazas (plazas asignadas a usuarios).
    """
    # Consulta para obtener abonados con nombre, NIF y plazas asociadas.
    sql = """
        SELECT
            CONCAT_WS(' ', pr.apellidos, pr.nombre) AS nombre_razon_social,
            pr.NIF                              AS nif,
            p.codigo_plazas                    AS codigo_plaza
        FROM parquin_camiones.tbl_historico_plazas AS h
        INNER JOIN parquin_camiones.tbl_plazas AS p
            ON h.idtbl_plazas = p.idtbl_plazas
        INNER JOIN bd_tbl_comunes.tbl_proveedores AS pr
            ON h.idtbl_proveedores = pr.Idtbl_proveedores
        WHERE
            h.fecha_inicio <= %s
            AND (h.fecha_fin IS NULL OR h.fecha_fin >= %s)
        ORDER BY pr.apellidos, pr.nombre, p.fila, p.codigo_plazas
    """
    filas = ejecutar_query(
        sql,
        params=(fecha_fin_periodo, fecha_inicio_periodo),
        nombre_bd="parquin_camiones",
    )

    # Construimos estructura:
    #   clave: (nombre_razon_social, nif)
    #   valor: lista de códigos de plaza
    datos: dict[tuple[str, str], list[str]] = {}
    for fila in filas:
        clave = (fila["nombre_razon_social"], fila["nif"])
        cod = fila["codigo_plaza"]
        datos.setdefault(clave, []).append(cod)

    return datos


def _obtener_resumen_plazas_rio_torio(
    fecha_inicio_periodo: date, fecha_fin_periodo: date
) -> dict:
    """
    Devuelve un resumen de plazas del parquin Rio Torio:
      - totales
      - libres
      - ocupadas

    Basado en:
      - BD: parquin_camiones
      - Tabla: tbl_plazas
      - Una plaza está ocupada si idtbl_usuarios IS NOT NULL.
    """
    sql = """
        SELECT
            COUNT(DISTINCT h.idtbl_plazas) AS total,
            0 AS libres,
            COUNT(DISTINCT h.idtbl_plazas) AS ocupadas
        FROM parquin_camiones.tbl_historico_plazas AS h
        WHERE
            h.fecha_inicio <= %s
            AND (h.fecha_fin IS NULL OR h.fecha_fin >= %s)
    """

    filas = ejecutar_query(
        sql,
        params=(fecha_fin_periodo, fecha_inicio_periodo),
        nombre_bd="parquin_camiones",
    )
    row = filas[0] if filas else {"total": 0, "libres": 0, "ocupadas": 0}

    return {
        "totales": row["total"],
        "libres": row["libres"],
        "ocupadas": row["ocupadas"],
    }


# =============================================================================
# 4️⃣ FUNCIÓN PRINCIPAL: GENERAR EL DOCUMENTO WORD (JUNTOS)
# =============================================================================


def _generar_informe_word_juntos(mes_informe: int, anio_informe: int) -> Path:
    """
    Genera el informe mensual de abonados del parking Río Torío (versión JUNTOS)
    para el MES/AÑO indicados, y lo guarda en un archivo .docx.

    Contenido del documento:
      1. Cabecera y texto introductorio del asunto y destinatario.
      2. Contadores por forma de pago (ENCIMA de la tabla):
          · total pago mensual = X
          · total pago trimestral = Y
          · total pago semestral = Z
          · total pago anual = W
          · (opcional) otros = O
      3. Tabla de abonados con:
          · Nombre / Razón social
          · NIF
          · Plazas (códigos separados por comas) en una celda
          · Nº total de plazas en la celda contigua
      4. Párrafo final y firma del Jefe de Policía Local.

    El fichero se guarda en RUTA_DESTINO con nombre:
      <Mes>_<Año>_Padron_de_parking_Rio_Torio_JUNTOS.docx
    """

    # -------------------------------------------------------------------------
    # 4.1 Cálculo del rango de fechas del mes consultado (informativo)
    # -------------------------------------------------------------------------
    fecha_inicio_periodo = date(anio_informe, mes_informe, 1)
    if mes_informe == 12:
        fecha_inicio_mes_siguiente = date(anio_informe + 1, 1, 1)
    else:
        fecha_inicio_mes_siguiente = date(anio_informe, mes_informe + 1, 1)
    fecha_fin_periodo = fecha_inicio_mes_siguiente - timedelta(days=1)

    nombre_mes = MESES_ES[mes_informe].capitalize()
    nombre_fichero = (
        f"{nombre_mes}_{anio_informe}_Padron_de_parking_Rio_Torio_JUNTOS.docx"
    )

    # -------------------------------------------------------------------------
    # 4.2 Preparación de la ruta de destino
    # -------------------------------------------------------------------------
    RUTA_DESTINO.mkdir(parents=True, exist_ok=True)
    ruta_completa = RUTA_DESTINO / nombre_fichero

    # -------------------------------------------------------------------------
    # 4.3 Creación del documento Word y cabeceras
    # -------------------------------------------------------------------------
    document = Document()
    section = document.sections[0]
    hoy = date.today()

    # Resumen de plazas para el encabezado (libres, ocupadas, totales)
    resumen_plazas = _obtener_resumen_plazas_rio_torio(
        fecha_inicio_periodo,
        fecha_fin_periodo,
    )
    totales = resumen_plazas["totales"]
    libres = resumen_plazas["libres"]
    ocupadas = resumen_plazas["ocupadas"]

    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_paragraph.add_run("Jefatura de Policía Local\nAdministración\n")
    header_paragraph.add_run(
        f"Plazas totales: {totales} | " f"Libres: {libres} | " f"Ocupadas: {ocupadas}"
    )

    section.left_margin = Cm(3)

    p_asunto = document.add_paragraph()
    p_asunto.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_asunto.add_run(
        "ASUNTO: Relación de abonados y plazas del "
        "parking de camiones de la Calle Río Torío (formato conjunto)."
    ).bold = True

    document.add_paragraph()
    p_interesado = document.add_paragraph()
    p_interesado.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_interesado.add_run("Interesado: Gestión Tributaria.")

    document.add_paragraph()
    p_comunico = document.add_paragraph()
    p_comunico.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_comunico.paragraph_format.first_line_indent = Cm(2)
    p_comunico.add_run(
        f"Comunico a Vd. que durante el mes de {nombre_mes} de {anio_informe} "
        "la relación de abonados y sus plazas es la siguiente:"
    )
    document.add_paragraph()

    # -------------------------------------------------------------------------
    # 4.4 Obtención de datos de abonados y plazas (versión juntos)
    # -------------------------------------------------------------------------
    abonados_plazas = _obtener_abonados_y_plazas_juntos(
        fecha_inicio_periodo,
        fecha_fin_periodo,
    )

    # -------------------------------------------------------------------------
    # 4.4.bis Contadores por forma de pago (ENCIMA de la tabla)
    # -------------------------------------------------------------------------
    contadores_fp = _contar_forma_pago_juntos(
        fecha_inicio_periodo,
        fecha_fin_periodo,
    )

    p_contadores = document.add_paragraph()
    p_contadores.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_contadores.paragraph_format.first_line_indent = Cm(2)

    texto_contadores = (
        f"total pago mensual = {contadores_fp['mensual']} | "
        f"trimestral = {contadores_fp['trimestral']} | "
        f"semestral = {contadores_fp['semestral']} | "
        f"anual = {contadores_fp['anual']}"
    )
    if contadores_fp["otros"]:
        texto_contadores += f" | otros = {contadores_fp['otros']}"

    p_contadores.add_run(texto_contadores).bold = True

    document.add_paragraph()

    # -------------------------------------------------------------------------
    # 4.5 Tabla de abonados con plazas en una celda y total en otra
    # -------------------------------------------------------------------------
    if abonados_plazas:
        # Tabla con cuatro columnas:
        #   1) Nombre / Razón social
        #   2) NIF
        #   3) Plazas (códigos separados por comas)
        #   4) Nº total de plazas
        table = document.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Nombre / Razón social"
        hdr_cells[1].text = "NIF"
        hdr_cells[2].text = "Plazas"
        hdr_cells[3].text = "Total plazas"

        for (nombre, nif), lista_plazas in abonados_plazas.items():
            row_cells = table.add_row().cells
            row_cells[0].text = str(nombre)
            row_cells[1].text = str(nif)
            # Códigos de plaza unidos por comas en una celda
            if lista_plazas:
                row_cells[2].text = ", ".join(str(c) for c in lista_plazas)
                row_cells[3].text = str(len(lista_plazas))
            else:
                row_cells[2].text = ""
                row_cells[3].text = "0"
    else:
        p_sin = document.add_paragraph()
        p_sin.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_sin.add_run(
            "Actualmente no existen abonados registrados con plazas asignadas."
        ).italic = True

    document.add_paragraph()

    # -------------------------------------------------------------------------
    # 4.6 Párrafo final y firma
    # -------------------------------------------------------------------------
    p_final = document.add_paragraph()
    p_final.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_final.paragraph_format.first_line_indent = Cm(2)
    p_final.add_run("Lo que comunico a Vd., para su conocimiento y efectos oportunos.")

    document.add_paragraph()
    p_firma_fecha = document.add_paragraph()
    p_firma_fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_firma_fecha.add_run(f"En Ávila, a {hoy.strftime('%d/%m/%Y')}.")

    document.add_paragraph()
    p_cargo = document.add_paragraph()
    p_cargo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cargo.add_run("El Jefe de la Policía Local.")

    document.add_paragraph()
    p_nombre = document.add_paragraph()
    p_nombre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_nombre.add_run("D. Carlos Blanco Rubio")

    p_digital = document.add_paragraph()
    p_digital.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_digital.add_run("(documento firmado digitalmente)").italic = True

    document.save(ruta_completa)
    return ruta_completa


# =============================================================================
# 5️⃣ HELPER: CONTAR FORMA DE PAGO (JUNTOS)
# =============================================================================


def _contar_forma_pago_juntos(
    fecha_inicio_periodo: date, fecha_fin_periodo: date
) -> dict:
    """
    Calcula contadores por forma de pago para el padrón JUNTOS.

    Aquí tienes dos opciones:
      1) Si ya tienes un SELECT histórico similar al de Río Torío técnico,
         llama a esa función y cuenta sobre sus filas.
      2) Si para JUNTOS usas la tabla de usuarios/proveedores directamente,
         haz un SELECT específico y cuenta aquí.

    Dejo un ejemplo genérico usando tbl_usuarios y tbl_plazas como fuente.
    """

    query = """
        SELECT
            h.forma_pago AS forma_pago_usuario,
            COUNT(DISTINCT h.idtbl_plazas) AS num_plazas
        FROM parquin_camiones.tbl_historico_plazas AS h
        WHERE
            h.fecha_inicio <= %s
            AND (h.fecha_fin IS NULL OR h.fecha_fin >= %s)
        GROUP BY h.forma_pago
    """

    filas = ejecutar_query(
        query,
        params=(fecha_fin_periodo, fecha_inicio_periodo),
        nombre_bd="parquin_camiones",
    )

    total_pago_mensual = 0
    total_pago_trimestral = 0
    total_pago_semestral = 0
    total_pago_anual = 0
    total_pago_otro = 0

    for f in filas:
        forma_raw = f.get("forma_pago_usuario") or ""
        forma = forma_raw.strip().lower()
        num_plazas = f.get("num_plazas", 0) or 0

        if forma in ("mensual", "m"):
            total_pago_mensual += num_plazas
        elif forma in ("trimestral", "t"):
            total_pago_trimestral += num_plazas
        elif forma in ("semestral", "s"):
            total_pago_semestral += num_plazas
        elif forma in ("anual", "a"):
            total_pago_anual += num_plazas
        else:
            total_pago_otro += num_plazas

    return {
        "mensual": total_pago_mensual,
        "trimestral": total_pago_trimestral,
        "semestral": total_pago_semestral,
        "anual": total_pago_anual,
        "otros": total_pago_otro,
    }


# =============================================================================
# 5️⃣ VISTA DEL BOTÓN: FORMULARIO (GET) + GENERACIÓN (POST) · JUNTOS
# =============================================================================


@btn_rio_torio_padron_juntos_bp.route(
    "/btn_rio_torio_padron_juntos_bp",
    methods=["GET", "POST"],
)
@rol_required("gestor", "super_admin")
def btn_rio_torio_padron_juntos():
    """
    Vista asociada al botón de generación del padrón mensual (Río Torío, versión
    JUNTOS: plazas en una celda y total en la celda contigua).

    Flujo:
      - Si el usuario no tiene sesión iniciada → redirige a login.
      - Si la petición es GET:
          · Calcula un mes y año por defecto (el mes anterior).
          · Renderiza el formulario para que el usuario elija MES y AÑO.
      - Si la petición es POST:
          · Valida los valores de MES y AÑO.
          · Llama a _generar_informe_word_juntos(mes, año).
          · Muestra un mensaje flash de éxito o error.
          · Devuelve el Word generado para su descarga/apertura.
    """
    if not session.get("user_id"):
        flash("Debes iniciar sesión", "danger")
        return redirect(url_for("auth_bp.login"))

    if request.method == "GET":
        hoy = date.today()
        if hoy.month == 1:
            mes_defecto = 12
            anio_defecto = hoy.year - 1
        else:
            mes_defecto = hoy.month - 1
            anio_defecto = hoy.year

        return render_template(
            "parquin/rio_torio/rio_torio_padron_juntos.html",
            meses=MESES_ES,
            mes_defecto=mes_defecto,
            anio_defecto=anio_defecto,
        )

    mes_str = request.form.get("mes")
    anio_str = request.form.get("anio")

    try:
        mes = int(mes_str)
        anio = int(anio_str)
        if mes < 1 or mes > 12:
            raise ValueError("Mes fuera de rango")
    except Exception:
        flash("Mes o año no válidos.", "danger")
        return redirect(
            url_for("btn_rio_torio_padron_juntos_bp.btn_rio_torio_padron_juntos")
        )

    try:
        ruta = _generar_informe_word_juntos(mes, anio)
        flash(f"Padrón (juntos) generado correctamente: {ruta.name}", "success")

        return send_file(
            ruta,
            as_attachment=True,
            download_name=ruta.name,
            mimetype=(
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            ),
        )
    except Exception as e:
        current_app.logger.error(f"Error generando padrón Rio Torío (juntos): {e}")
        flash("Error generando el informe de padrón (juntos).", "danger")
        return redirect(
            url_for("btn_rio_torio_padron_juntos_bp.btn_rio_torio_padron_juntos")
        )

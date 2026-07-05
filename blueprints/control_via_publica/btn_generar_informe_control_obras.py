# =============================================================================
# 🧾 BOTÓN OCUPACIÓN VÍA · INFORME MENSUAL CONTROL DE OBRAS
# =============================================================================
# 🔍 INTRODUCCIÓN GENERAL
# -----------------------------------------------------------------------------
# Este código convierte el script standalone de generación de informes
# "generar_informe_control_obras.py" en un BOTÓN Flask:
#
#   btn_ocupacion_via_informe_mensual
#
# que cuelga del módulo de OCUPACIÓN DE VÍA PÚBLICA.
#
# 📌 ¿QUÉ HACE ESTE BOTÓN?
#   1) Calcula el PERIODO DEL MES ANTERIOR:
#        - desde el día 1 del mes anterior
#        - hasta el día 1 del mes actual (intervalo [inicio, fin))
#   2) Consulta en BD:
#        - TODAS las obras SIN fecha_finalizacion (obras abiertas)
#        - TODAS las inspecciones (controles) de ese mes anterior
#   3) Agrupa los datos:
#        - Una entrada por obra
#        - Lista de visitas/controles asociados
#        - Si no hay visitas, la obra aparece igualmente con "Sin visitas"
#   4) Genera:
#        - Un DOCX (python-docx) con una sección por obra
#        - Un PDF (reportlab) con el resumen de las mismas obras
#   5) Guarda auditoría en BD:
#        - Rutas a DOCX y PDF
#        - Estado (ok / error)
#        - Detalle de error si lo hay
#
# 📁 Las rutas de salida se leen desde configuración:
#
#   bd_tbl_comunes.tbl_app_config
#   clave = 'RUTA_REPORTES_CONTROL_OBRAS'
#
# 🧩 Transformación a botón:
#   - Se añade una ruta Flask:
#         /ocupacion_via/btn_ocupacion_via_informe_mensual
#   - Protegida con login / rol
#   - Al pulsar el botón:
#         - genera DOCX + PDF
#         - registra auditoría
#         - devuelve un mensaje (flash / plantilla)
# =============================================================================


# =============================================================================
# 1️⃣ IMPORTACIONES GENERALES
# =============================================================================
# 1.1) Tipos y fechas
# -----------------------------------------------------------------------------
from __future__ import annotations

import os
from dataclasses import dataclass  # (no se usa aquí directamente, pero se mantiene por si amplías)
from datetime import date, datetime
from typing import Dict, Any, List, Tuple

# 1.2) Librerías de generación de documentos
# -----------------------------------------------------------------------------
from docx import Document                    # 📄 Word (python-docx)
from reportlab.lib.pagesizes import A4       # 📄 Tamaño de página para PDF
from reportlab.pdfgen import canvas          # 🖨�?Generación de PDF

# 1.3) Flask y seguridad
# -----------------------------------------------------------------------------
from flask import Blueprint, render_template, current_app, redirect, url_for, flash
from flask_login import login_required
from services.helpers import rol_required

# 1.4) Acceso a base de datos
# -----------------------------------------------------------------------------
from db import ejecutar_query, ejecutar_non_query
# =============================================================================
# 1️⃣ FIN IMPORTACIONES
# =============================================================================



# =============================================================================
# 2️⃣ BLUEPRINT DEL MÓDULO OCUPACIÓN VÍA · BOTÓN INFORME
# =============================================================================
# Este blueprint representa el MÓDULO de OCUPACIÓN DE VÍA PÚBLICA (control de obras).
# De él colgará el botón:
#
#   btn_ocupacion_via_informe_mensual
#
# URL base del módulo (ajusta según tu estructura real):
#   /ocupacion_via
# =============================================================================
modulo_ocupacion_via_bp = Blueprint(
    "modulo_ocupacion_via_bp",
    __name__,
    url_prefix="/ocupacion_via",
)
# =============================================================================
# 2️⃣ FIN BLUEPRINT
# =============================================================================



# =============================================================================
# 3️⃣ CÁLCULO DEL PERIODO · RANGO MES ANTERIOR
# =============================================================================
# 3.1) Función rango_mes_anterior
# -----------------------------------------------------------------------------
#   - Entrada: fecha opcional (por defecto, hoy)
#   - Salida:
#       * inicio (YYYY-MM-01 del mes anterior)
#       * fin    (YYYY-MM-01 del mes actual)
#       * año del mes anterior
#       * mes del mes anterior
#
#   Ejemplo:
#       hoy = 2026-04-10
#       �?inicio = 2026-03-01
#       �?fin    = 2026-04-01
#       �?anio   = 2026
#       �?mes    = 3
# =============================================================================
def rango_mes_anterior(hoy: date | None = None) -> Tuple[date, date, int, int]:
    """
    Devuelve:
      - inicio (YYYY-MM-01 del mes anterior)
      - fin    (YYYY-MM-01 del mes actual)
      - año del mes anterior
      - mes del mes anterior
    """
    hoy = hoy or date.today()
    year = hoy.year
    month = hoy.month

    # 🧮 mes anterior
    if month == 1:
        prev_year = year - 1
        prev_month = 12
    else:
        prev_year = year
        prev_month = month - 1

    inicio = date(prev_year, prev_month, 1)
    fin = date(year, month, 1)
    return inicio, fin, prev_year, prev_month
# =============================================================================
# 3️⃣ FIN CÁLCULO PERIODO
# =============================================================================



# =============================================================================
# 4️⃣ LECTURA CARPETA DESTINO DESDE CONFIG
# =============================================================================
# 4.1) get_ruta_reportes()
# -----------------------------------------------------------------------------
#   - Lee la ruta base donde se guardan los informes desde:
#
#       bd_tbl_comunes.tbl_app_config
#       clave = 'RUTA_REPORTES_CONTROL_OBRAS'
#
#   - Si no existe esa clave �?lanza RuntimeError
# =============================================================================
def get_ruta_reportes() -> str:
    sql = """
      SELECT valor
      FROM bd_tbl_comunes.tbl_app_config
      WHERE clave = 'RUTA_REPORTES_CONTROL_OBRAS'
      LIMIT 1
    """
    filas = ejecutar_query(sql, params=(), nombre_bd="control_via_publica")
    if not filas or not filas[0].get("valor"):
        raise RuntimeError(
            "No existe la config RUTA_REPORTES_CONTROL_OBRAS en bd_tbl_comunes.tbl_app_config"
        )
    return filas[0]["valor"]
# =============================================================================
# 4️⃣ FIN LECTURA CARPETA DESTINO
# =============================================================================



# =============================================================================
# 5️⃣ QUERY PRINCIPAL · OBRAS ABIERTAS + CONTROLES DEL MES
# =============================================================================
# 5.1) cargar_datos_informe(inicio, fin)
# -----------------------------------------------------------------------------
#   - Recupera:
#       * Obras sin fecha_finalizacion (abiertas)
#       * Controles (tbl_control_via_publica) entre [inicio, fin)
#   - Incluye información de:
#       * proveedor (nombre, NIF)
#       * tipo de vía, calle
#       * detalles de cada inspección (vallas, materiales, andamios, grúas...)
# =============================================================================
def cargar_datos_informe(inicio: date, fin: date) -> List[Dict[str, Any]]:
    sql = """
    SELECT
      o.idtbl_obras,
      o.idtbl_gis_municipal,
      o.idtbl_proveedor,
      o.idtbl_municipios,
      o.idtbl_tipos_de_vias,
      o.idtbl_calles,
      o.observaciones,
      o.lat,
      o.lon,
      o.idtbl_gestores,
      o.fecha_inicio,
      o.fecha_finalizacion,

      p.Nombre_Razon_Social AS proveedor_nombre,
      p.NIF AS proveedor_nif,

      c.calles AS calle_nombre,
      tv.tipos_de_vias AS tipo_via_nombre,

      cvp.idtbl_control_via_publica,
      cvp.fecha_inspeccion,
      cvp.vallas,
      cvp.vallas_metros,
      cvp.materiales_de_construccion,
      cvp.materiales_metros,
      cvp.silos,
      cvp.silos_metros,
      cvp.andamios,
      cvp.andamios_metros,
      cvp.gruas,
      cvp.gruas_metros,
      cvp.n_agente1,
      cvp.idtbl_gestores AS gestor_control_id

    FROM control_via_publica.tbl_obras o

    LEFT JOIN bd_tbl_comunes.tbl_proveedores p
      ON p.Idtbl_proveedores = o.idtbl_proveedor

    LEFT JOIN bd_tbl_comunes.tbl_calles c
      ON c.idtbl_calles = o.idtbl_calles

    LEFT JOIN bd_tbl_comunes.tbl_tipos_de_vias tv
      ON tv.idtbl_tipos_de_vias = o.idtbl_tipos_de_vias

    LEFT JOIN control_via_publica.tbl_control_via_publica cvp
      ON cvp.idtbl_obras = o.idtbl_obras
     AND cvp.fecha_inspeccion >= %s
     AND cvp.fecha_inspeccion <  %s

    WHERE (o.fecha_finalizacion IS NULL OR o.fecha_finalizacion = '')

    ORDER BY o.idtbl_obras DESC, cvp.fecha_inspeccion ASC
    """
    return ejecutar_query(sql, params=(inicio, fin), nombre_bd="control_via_publica")
# =============================================================================
# 5️⃣ FIN QUERY PRINCIPAL
# =============================================================================



# =============================================================================
# 6️⃣ AGRUPAR POR OBRA · ESTRUCTURA PARA INFORME
# =============================================================================
# 6.1) agrupar_por_obra(filas)
# -----------------------------------------------------------------------------
#   - Entrada: lista de filas de BD
#   - Salida:
#       {
#         idobra: {
#           "obra": { ...datos obra... },
#           "visitas": [
#               { ...datos inspección 1... },
#               { ...datos inspección 2... },
#               ...
#           ]
#         },
#         ...
#       }
#
#   - Si una obra no tiene visitas en el mes:
#       "visitas" será lista vacía �?se mostrará "Sin visitas"
# =============================================================================
def agrupar_por_obra(filas: List[Dict[str, Any]]) -> Dict[int, Dict[str, Any]]:
    """
    Devuelve:
      { idobra: { 'obra': {...}, 'visitas': [ {...}, {...} ] } }
    """
    out: Dict[int, Dict[str, Any]] = {}

    for r in filas:
        oid = r["idtbl_obras"]
        if oid not in out:
            out[oid] = {
                "obra": {
                    "idtbl_obras": r["idtbl_obras"],
                    "proveedor": r.get("proveedor_nombre"),
                    "proveedor_nif": r.get("proveedor_nif"),
                    "tipo_via": r.get("tipo_via_nombre"),
                    "calle": r.get("calle_nombre"),
                    "observaciones": r.get("observaciones"),
                    "fecha_inicio": r.get("fecha_inicio"),
                    "fecha_finalizacion": r.get("fecha_finalizacion"),
                },
                "visitas": [],
            }

        # �?Si hay control ese mes, se añade a visitas
        if r.get("idtbl_control_via_publica"):
            out[oid]["visitas"].append({
                "fecha_inspeccion": r.get("fecha_inspeccion"),
                "n_agente1": r.get("n_agente1"),
                "vallas": r.get("vallas"),
                "vallas_metros": r.get("vallas_metros"),
                "materiales_de_construccion": r.get("materiales_de_construccion"),
                "materiales_metros": r.get("materiales_metros"),
                "silos": r.get("silos"),
                "silos_metros": r.get("silos_metros"),
                "andamios": r.get("andamios"),
                "andamios_metros": r.get("andamios_metros"),
                "gruas": r.get("gruas"),
                "gruas_metros": r.get("gruas_metros"),
                "gestor_control_id": r.get("gestor_control_id"),
            })

    return out
# =============================================================================
# 6️⃣ FIN AGRUPACIÓN
# =============================================================================



# =============================================================================
# 7️⃣ GENERAR DOCX · INFORME DETALLADO
# =============================================================================
# 7.1) generar_docx(path_docx, datos, periodo_anio, periodo_mes)
# -----------------------------------------------------------------------------
#   - Crea un documento Word (.docx)
#   - Añade cabecera con título y fecha de generación
#   - Por cada obra:
#       * Muestra datos básicos
#       * Genera tabla con las visitas del mes
#       * O indica "Sin visitas"
# =============================================================================
def generar_docx(path_docx: str, datos: Dict[int, Dict[str, Any]], periodo_anio: int, periodo_mes: int) -> None:
    doc = Document()

    # 🏷�?Título
    doc.add_heading(
        f"Informe mensual Control de Obras - {periodo_anio}-{periodo_mes:02d}",
        level=1,
    )
    doc.add_paragraph(
        f"Fecha de generación: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )

    doc.add_paragraph(
        "Incluye todas las obras SIN fecha de finalización. "
        "Cada obra muestra las inspecciones realizadas en el periodo. "
        "Si no hay inspecciones, se indica 'Sin visitas'."
    )

    # 🔁 Una sección por obra
    for obra_id, pack in datos.items():
        obra = pack["obra"]
        visitas = pack["visitas"]

        doc.add_heading(f"Obra #{obra_id}", level=2)

        # 📌 Datos básicos obra
        doc.add_paragraph(
            f"Proveedor: {obra.get('proveedor') or '-'} "
            f"({obra.get('proveedor_nif') or '-'})"
        )
        doc.add_paragraph(
            f"Ubicación: {obra.get('tipo_via') or '-'} "
            f"{obra.get('calle') or '-'}"
        )
        doc.add_paragraph(
            f"Observaciones: {obra.get('observaciones') or '-'}"
        )

        # 📋 Tabla de visitas (o mensaje)
        if not visitas:
            doc.add_paragraph("🟡 Sin visitas registradas en este mes.")
        else:
            table = doc.add_table(rows=1, cols=9)
            hdr = table.rows[0].cells
            hdr[0].text = "Fecha"
            hdr[1].text = "Agente"
            hdr[2].text = "Vallas"
            hdr[3].text = "m"
            hdr[4].text = "Materiales"
            hdr[5].text = "m"
            hdr[6].text = "Andamios"
            hdr[7].text = "Grúas"
            hdr[8].text = "GestorID"

            for v in visitas:
                row = table.add_row().cells
                row[0].text = str(v.get("fecha_inspeccion") or "-")
                row[1].text = str(v.get("n_agente1") or "-")
                row[2].text = "Sí" if v.get("vallas") else "No"
                row[3].text = str(v.get("vallas_metros") or "-")
                row[4].text = "Sí" if v.get("materiales_de_construccion") else "No"
                row[5].text = str(v.get("materiales_metros") or "-")
                row[6].text = "Sí" if v.get("andamios") else "No"
                row[7].text = "Sí" if v.get("gruas") else "No"
                row[8].text = str(v.get("gestor_control_id") or "-")

        doc.add_paragraph("�? * 40)

    # 💾 Guardar
    os.makedirs(os.path.dirname(path_docx), exist_ok=True)
    doc.save(path_docx)
# =============================================================================
# 7️⃣ FIN GENERACIÓN DOCX
# =============================================================================



# =============================================================================
# 8️⃣ GENERAR PDF · RESUMEN PDF DEL INFORME
# =============================================================================
# 8.1) generar_pdf(path_pdf, datos, periodo_anio, periodo_mes)
# -----------------------------------------------------------------------------
#   - Crea un PDF con reportlab
#   - Recorre las obras y sus visitas
#   - Hace salto de página cuando no hay espacio
# =============================================================================
def generar_pdf(path_pdf: str, datos: Dict[int, Dict[str, Any]], periodo_anio: int, periodo_mes: int) -> None:
    os.makedirs(os.path.dirname(path_pdf), exist_ok=True)

    c = canvas.Canvas(path_pdf, pagesize=A4)
    width, height = A4

    y = height - 50
    c.setFont("Helvetica-Bold", 14)
    c.drawString(
        50,
        y,
        f"Informe mensual Control de Obras - {periodo_anio}-{periodo_mes:02d}",
    )
    y -= 20
    c.setFont("Helvetica", 10)
    c.drawString(50, y, f"Generado: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    y -= 30

    for obra_id, pack in datos.items():
        obra = pack["obra"]
        visitas = pack["visitas"]

        if y < 140:
            c.showPage()
            y = height - 50

        c.setFont("Helvetica-Bold", 12)
        c.drawString(50, y, f"Obra #{obra_id}")
        y -= 16

        c.setFont("Helvetica", 10)
        c.drawString(
            50,
            y,
            f"Proveedor: {obra.get('proveedor') or '-'} "
            f"({obra.get('proveedor_nif') or '-'})",
        )
        y -= 14
        c.drawString(
            50,
            y,
            f"Ubicación: {obra.get('tipo_via') or '-'} "
            f"{obra.get('calle') or '-'}",
        )
        y -= 14
        obs = (obra.get("observaciones") or "-")[:120]
        c.drawString(50, y, f"Obs: {obs}")
        y -= 16

        if not visitas:
            c.drawString(50, y, "🟡 Sin visitas registradas en este mes.")
            y -= 18
        else:
            # cabecera simple
            c.setFont("Helvetica-Bold", 9)
            c.drawString(50, y, "Fecha")
            c.drawString(120, y, "Agente")
            c.drawString(200, y, "Vallas(m)")
            c.drawString(280, y, "Mat(m)")
            c.drawString(350, y, "And")
            c.drawString(390, y, "Grúas")
            c.drawString(440, y, "GestorID")
            y -= 12
            c.setFont("Helvetica", 9)

            for v in visitas:
                if y < 90:
                    c.showPage()
                    y = height - 50
                c.drawString(50, y, str(v.get("fecha_inspeccion") or "-")[:10])
                c.drawString(120, y, str(v.get("n_agente1") or "-")[:10])
                c.drawString(
                    200,
                    y,
                    f"{'Sí' if v.get('vallas') else 'No'}({v.get('vallas_metros') or '-'})",
                )
                c.drawString(
                    280,
                    y,
                    f"{'Sí' if v.get('materiales_de_construccion') else 'No'}"
                    f"({v.get('materiales_metros') or '-'})",
                )
                c.drawString(350, y, "Sí" if v.get("andamios") else "No")
                c.drawString(390, y, "Sí" if v.get("gruas") else "No")
                c.drawString(440, y, str(v.get("gestor_control_id") or "-"))
                y -= 12

            y -= 8

        c.setFont("Helvetica", 10)
        c.drawString(50, y, "-" * 80)
        y -= 18

    c.save()
# =============================================================================
# 8️⃣ FIN GENERACIÓN PDF
# =============================================================================



# =============================================================================
# 9️⃣ REGISTRO DE AUDITORÍA EN BD
# =============================================================================
# 9.1) registrar_informe(...)
# -----------------------------------------------------------------------------
#   - Inserta una fila en control_via_publica.tbl_informes_control_obras
#   - Guarda:
#       * periodo_anio, periodo_mes
#       * ruta_docx, ruta_pdf
#       * estado ("ok"/"error")
#       * error_detalle (texto)
# =============================================================================
def registrar_informe(
    periodo_anio: int,
    periodo_mes: int,
    ruta_docx: str,
    ruta_pdf: str,
    estado: str,
    error: str | None,
) -> None:
    sql = """
      INSERT INTO control_via_publica.tbl_informes_control_obras
        (periodo_anio, periodo_mes, ruta_docx, ruta_pdf, estado, error_detalle)
      VALUES (%s, %s, %s, %s, %s, %s)
    """
    ejecutar_non_query(
        sql,
        params=(periodo_anio, periodo_mes, ruta_docx, ruta_pdf, estado, error),
        nombre_bd="control_via_publica",
    )
# =============================================================================
# 9️⃣ FIN REGISTRO AUDITORÍA
# =============================================================================



# =============================================================================
# 🔟 BOTÓN FLASK · btn_ocupacion_via_informe_mensual
# =============================================================================
# 10.1) Ruta del botón
# -----------------------------------------------------------------------------
#   - URL:
#       /ocupacion_via/btn_ocupacion_via_informe_mensual
#
#   - Acceso:
#       - login_required
#       - rol_required("super_admin")
#
#   - Flujo:
#       1) Calcula rango del mes anterior
#       2) Obtiene carpeta base de reportes
#       3) Carga datos desde BD
#       4) Agrupa por obra
#       5) Genera DOCX + PDF
#       6) Registra auditoría
#       7) Muestra mensaje flash y redirige al módulo de ocupación de vía
# =============================================================================
@modulo_ocupacion_via_bp.route(
    "/btn_ocupacion_via_informe_mensual",
    methods=["GET"],
)
@login_required
@rol_required("super_admin")
def btn_ocupacion_via_informe_mensual():
    """
    BOTÓN · Genera informe mensual de control de obras (ocupación de vía).
    """
    try:
        # 1) Calcular periodo
        inicio, fin, anio, mes = rango_mes_anterior()

        # 2) Resolver ruta de reportes base
        carpeta_base = get_ruta_reportes()
        os.makedirs(carpeta_base, exist_ok=True)

        nombre_base = f"control_obras_{anio}_{mes:02d}"
        ruta_docx = os.path.join(carpeta_base, f"{nombre_base}.docx")
        ruta_pdf = os.path.join(carpeta_base, f"{nombre_base}.pdf")

        # 3) Cargar datos de BD
        filas = cargar_datos_informe(inicio, fin)
        datos = agrupar_por_obra(filas)

        # 4) Generar DOCX y PDF
        generar_docx(ruta_docx, datos, anio, mes)
        generar_pdf(ruta_pdf, datos, anio, mes)

        # 5) Registrar auditoría
        registrar_informe(anio, mes, ruta_docx, ruta_pdf, "ok", None)

        flash(
            f"Informe mensual de ocupación de vía generado correctamente "
            f"({ruta_docx}, {ruta_pdf}).",
            "success",
        )
    except Exception as e:
        # Registrar error en auditoría si es posible
        try:
            # Si algo ha fallado antes de definir rutas, las reconstruimos mínimamente
            inicio, fin, anio, mes = rango_mes_anterior()
            carpeta_base = get_ruta_reportes()
            nombre_base = f"control_obras_{anio}_{mes:02d}"
            ruta_docx = os.path.join(carpeta_base, f"{nombre_base}.docx")
            ruta_pdf = os.path.join(carpeta_base, f"{nombre_base}.pdf")
            registrar_informe(anio, mes, ruta_docx, ruta_pdf, "error", str(e))
        except Exception:
            pass

        current_app.logger.error(
            f"Error generando informe mensual ocupación vía: {e}"
        )
        flash(
            "Error generando el informe mensual de ocupación de vía.",
            "danger",
        )

    # Redirigir al módulo principal de ocupación de vía (ajusta nombre de vista)
    return redirect(
        url_for("modulo_ocupacion_via_bp.modulo_ocupacion_via")
    )
# =============================================================================
# 🔟 FIN BOTÓN FLASK · btn_ocupacion_via_informe_mensual
# =============================================================================